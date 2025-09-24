from flask import Flask, render_template, request, jsonify, send_from_directory, redirect
from flask_cors import CORS
import google.generativeai as genai
from dotenv import load_dotenv
import os
import logging
import json
import time
import uuid
from datetime import datetime, timezone
import jwt
import hashlib
import random
import secrets
import requests
from urllib.parse import urlencode
import re
import io
import csv

# Firebase Admin SDK
try:
    import firebase_admin
    from firebase_admin import credentials as fb_credentials
    from firebase_admin import db as fb_db
except Exception:
    firebase_admin = None
    fb_credentials = None
    fb_db = None

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables from multiple likely locations
# 1) current working directory, 2) this file's directory
try:
    # Load default if present in CWD
    load_dotenv()
    # Load from this module directory
    _this_dir = os.path.dirname(__file__)
    _this_env = os.path.join(_this_dir, '.env')
    if os.path.exists(_this_env):
        load_dotenv(_this_env, override=False)
except Exception:
    pass

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Initialize Firebase (Realtime Database) for appointments
rtdb_available = False
try:
    if firebase_admin is not None:
        # Get database URL from environment or use your Firebase project
        rtdb_url = os.getenv('FIREBASE_DB_URL', "https://livecode-35eda-default-rtdb.firebaseio.com/")

        try:
            firebase_admin.get_app()
        except ValueError:
            cred_path = os.getenv('GOOGLE_APPLICATION_CREDENTIALS')
            base_dir = os.path.dirname(__file__)
            # Resolve absolute and local file options
            candidate_paths = []
            # Root-level credentials.json (new key)
            candidate_paths.append(os.path.join(base_dir, 'livecode-35eda-firebase-adminsdk-fbsvc-c1c1d124c8.json'))
            if cred_path:
                candidate_paths.append(cred_path)
                candidate_paths.append(os.path.join(base_dir, cred_path))
            # Preferred new service account file (user-provided)
            candidate_paths.append(os.path.join(base_dir, 'livecode-35eda-firebase-adminsdk-fbsvc-4549eea2ad.json'))
            # Newly generated service account file (user-provided)
            candidate_paths.append(os.path.join(base_dir, 'livecode-35eda-firebase-adminsdk-fbsvc-a52c90eca7.json'))
            

            cred = None
            for cpath in candidate_paths:
                try:
                    if cpath and os.path.exists(cpath):
                        cred = fb_credentials.Certificate(cpath)
                        break
                except Exception:
                    pass
            if cred is None:
                cred = fb_credentials.ApplicationDefault()
            firebase_admin.initialize_app(cred, options={'databaseURL': rtdb_url})
        
        # Check RTDB availability
        try:
            _ = fb_db.reference('/')
            rtdb_available = True
        except Exception as e:
            logger.warning(f"Realtime Database not available: {e}")
    else:
        logger.warning('firebase_admin is not installed. Appointment storage will be disabled.')
except Exception as e:
    logger.error(f"Firebase init error: {e}")

# Appointment state management
appointment_state = {}

# Lead capture state management
lead_capture_state = {}  # {username_session: {'step': 1, 'name': '', 'email': '', 'phone': ''}}
message_counters = {}  # {username_session: count}
lead_capture_completed = {}  # {username_session: True/False}

 # Configure Gemini
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY', 'AIzaSyCfFyHf5i6lizMSFVdJXwxgCJH_rTuYTkY')
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-2.0-flash')
JWT_SECRET = os.getenv('JWT_SECRET', 'dev-secret-change')
JWT_EXP_SECONDS = 60 * 60 * 24 * 7


# In-memory cache per user for knowledge base
# Global knowledge base cache (cleared)
KB_CACHE = {}
# --- Phone normalization helper ---
def _normalize_phone(raw: str) -> str:
    try:
        raw = (raw or '').strip()
        if not raw:
            return ''
        digits = re.sub(r'[^\d]', '', raw)
        # If starts with country code 91 and total 12 digits, take last 10
        if len(digits) >= 12 and digits.startswith('91'):
            digits_last10 = digits[-10:]
            return '+91' + digits_last10
        # If exactly 10 digits, assume India and prefix +91
        if len(digits) == 10:
            return '+91' + digits
        # If user included country code with + and reasonable length
        if raw.startswith('+') and 10 <= len(digits) <= 15:
            return '+' + digits
        # Fallback: if 11-15 digits, still return with '+'
        if 11 <= len(digits) <= 15:
            return '+' + digits
        return ''
    except Exception:
        return ''

# Cache for Gemini responses
response_cache = {}
CACHE_EXPIRY = 3600  # Cache expiry time in seconds (1 hour)

def get_cached_response(user_input):
    """Get cached response if available and not expired"""
    current_time = time.time()
    if user_input in response_cache:
        cached_time, cached_response = response_cache[user_input]
        if current_time - cached_time < CACHE_EXPIRY:
            return cached_response
    return None

def cache_response(user_input, response):
    """Cache the response with current timestamp"""
    response_cache[user_input] = (time.time(), response)

def load_knowledge_base_for(username: str):
    """Load knowledge base for a specific user and cache it."""
    global COMMON_QUESTIONS
    try:
        if not rtdb_available or not username:
            return KB_CACHE.get(username or '', {})
        knowledge_ref = fb_db.reference(f'{username}/knowledge_base')
        knowledge_data = knowledge_ref.get() or {}
        KB_CACHE[username] = knowledge_data
        if isinstance(knowledge_data, dict) and 'common_questions' in knowledge_data:
            COMMON_QUESTIONS = knowledge_data.get('common_questions') or {}
        try:
            response_cache.clear()
        except Exception:
            pass
        return knowledge_data
    except Exception as e:
        logger.error(f"Error loading knowledge base for {username}: {e}")
        return KB_CACHE.get(username or '', {})

def format_knowledge_base(kb_data):
    """Convert knowledge base data to a readable format, filtering out technical details"""
    if not kb_data:
        return ""
    
    context_parts = []
    
    # Technical fields to skip (more specific to avoid filtering business data)
    technical_fields = {
        'password', 'last_login', 'created_at', 'updated_at', 
        'user_id', 'session_id', 'token', 'key', 'secret', 'hash', 'salt',
        'database', 'table', 'admin_id', 'login_time'
    }
    
    def is_technical_key(key):
        key_lower = str(key).lower()
        return any(tech in key_lower for tech in technical_fields)
    
    def process_value(key, value, indent=""):
        # Skip technical fields
        if is_technical_key(key):
            return
            
        if isinstance(value, dict):
            # Only process if it contains non-technical data
            has_relevant_data = any(not is_technical_key(k) for k in value.keys())
            if has_relevant_data:
                for k, v in value.items():
                    if not is_technical_key(k):
                        process_value(k, v, indent + "- ")
        elif isinstance(value, list):
            # Filter out technical items from lists
            relevant_items = []
            for item in value:
                if isinstance(item, dict):
                    # Only include if it has non-technical fields
                    if any(not is_technical_key(k) for k in item.keys()):
                        relevant_items.append(item)
                elif not is_technical_key(str(item)):
                    relevant_items.append(item)
            
            if relevant_items:
                if len(relevant_items) <= 5:
                    context_parts.append(f"{indent}{key}: {', '.join(map(str, relevant_items))}")
                else:
                    context_parts.append(f"{indent}{key}: {', '.join(map(str, relevant_items[:3]))} and {len(relevant_items)-3} more")
        else:
            # Only include non-technical string values
            if not is_technical_key(str(value)) and len(str(value).strip()) > 0:
                context_parts.append(f"{indent}{key}: {str(value)}")
    
    # Process the knowledge base data
    if isinstance(kb_data, dict):
        for key, value in kb_data.items():
            if not is_technical_key(key):
                process_value(key, value)
    elif isinstance(kb_data, list):
        context_parts.append("Available information:")
        for i, item in enumerate(kb_data[:10]):  # Limit to first 10 items
            if isinstance(item, dict):
                # Only process if it has non-technical fields
                if any(not is_technical_key(k) for k in item.keys()):
                    for k, v in item.items():
                        if not is_technical_key(k):
                            process_value(k, v, "- ")
            elif not is_technical_key(str(item)):
                context_parts.append(f"- Item {i+1}: {str(item)}")
    else:
        if not is_technical_key(str(kb_data)):
            context_parts.append(f"Information: {str(kb_data)}")
    
    result = "\n".join(context_parts) if context_parts else ""
    
    # If no content was found, show a sample of the raw data for debugging
    if not result and kb_data:
        logger.warning("No business content found after filtering, showing raw data sample")
        if isinstance(kb_data, dict):
            sample_keys = list(kb_data.keys())[:5]
            result = f"Available data keys: {', '.join(sample_keys)}"
        elif isinstance(kb_data, list) and len(kb_data) > 0:
            result = f"Available data: {len(kb_data)} items"
    
    return result

def update_knowledge_base(new_data, username: str):
    """Update knowledge base data in Firebase for a user"""
    try:
        if rtdb_available and username:
            knowledge_ref = fb_db.reference(f'{username}/knowledge_base')
            knowledge_ref.set(new_data)
            KB_CACHE[username] = new_data
            try:
                response_cache.clear()
            except Exception:
                pass
            return True
        else:
            logger.warning("Database not available or username missing; cannot update knowledge base")
            return False
    except Exception as e:
        logger.error(f"Error updating knowledge base: {e}")
        return False

# --- Utility: compute per-user/per-bot base path ---
def _base_path(username: str, bot_id: str = '') -> str:
    username = (username or '').strip() or 'anonymous'
    bot_id = (bot_id or '').strip()
    if bot_id:
        return f"{username}/bots/{bot_id}"
    return username

# Common questions and their responses (will be populated from knowledge base)
COMMON_QUESTIONS = {}

# --- Password hashing functions ---
def _hash_password(password: str) -> tuple:
    """Hash password with salt and return (salt, hash)"""
    salt = secrets.token_hex(16)
    password_hash = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt.encode('utf-8'), 100000)
    return salt, password_hash.hex()

def _verify_password(password: str, salt: str, stored_hash: str) -> bool:
    """Verify password against stored hash"""
    password_hash = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt.encode('utf-8'), 100000)
    return password_hash.hex() == stored_hash

# --- Authentication middleware ---
def require_auth(f):
    """Decorator to require authentication for routes"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Try to get token from various sources
        token = None
        
        # 1. Check Authorization header
        auth_header = request.headers.get('Authorization', '')
        if auth_header.startswith('Bearer '):
            token = auth_header.replace('Bearer ', '')
        
        # 2. Check request JSON body
        if not token and request.is_json:
            token = request.json.get('token')
        
        # 3. Check query parameters
        if not token:
            token = request.args.get('token')
        
        # 4. Check cookies
        if not token:
            token = request.cookies.get('token')
        
        if not token:
            return jsonify({'success': False, 'message': 'Authentication required'}), 401
        
        username = _verify_token(token)
        if not username:
            return jsonify({'success': False, 'message': 'Invalid token'}), 401
        
        request.current_user = username
        return f(*args, **kwargs)
    return decorated_function

def require_auth_lenient(f):
    """Lenient auth decorator - allows requests without authentication for testing"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Try to get token from various sources
        token = None
        
        # 1. Check Authorization header
        auth_header = request.headers.get('Authorization', '')
        if auth_header.startswith('Bearer '):
            token = auth_header.replace('Bearer ', '')
        
        # 2. Check request JSON body
        if not token and request.is_json:
            token = request.json.get('token')
        
        # 3. Check query parameters
        if not token:
            token = request.args.get('token')
        
        # 4. Check cookies
        if not token:
            token = request.cookies.get('token')
        
        # For development/testing, allow requests without token
        if token:
            username = _verify_token(token)
            if username:
                request.current_user = username
            else:
                # Invalid token, unauthenticated
                request.current_user = ''
        else:
            # No token provided
            request.current_user = ''
        
        return f(*args, **kwargs)
    return decorated_function

def get_company_config(username):
    """Get company configuration for a user"""
    try:
        if not rtdb_available or not username:
            return None
        
        config_ref = fb_db.reference(f'{username}/company_config')
        config = config_ref.get()
        return config
    except Exception as e:
        logger.error(f"Error loading company config for {username}: {e}")
        return None

def get_personality_instructions(tone, industry):
    """Get detailed personality instructions based on tone and industry"""
    
    # Industry-specific adaptations
    industry_context = ""
    if industry.lower() in ['healthcare', 'medical', 'hospital', 'clinic', 'wellness']:
        industry_context = "IMPORTANT: This is a healthcare/medical context. Always maintain a caring, professional, and empathetic tone. Avoid humor that might be inappropriate. Focus on patient care and safety."
    elif industry.lower() in ['legal', 'law', 'attorney', 'lawyer']:
        industry_context = "IMPORTANT: This is a legal context. Maintain a formal, respectful, and professional tone. Avoid casual language or humor. Be precise and accurate."
    elif industry.lower() in ['finance', 'banking', 'investment', 'insurance']:
        industry_context = "IMPORTANT: This is a financial context. Be professional, trustworthy, and precise. Avoid casual language. Focus on accuracy and security."
    elif industry.lower() in ['entertainment', 'gaming', 'media', 'sports']:
        industry_context = "IMPORTANT: This is an entertainment context. You can be more casual, fun, and engaging. Use appropriate humor and emojis."
    elif industry.lower() in ['restaurant', 'food', 'hospitality', 'tourism']:
        industry_context = "IMPORTANT: This is a hospitality context. Be warm, welcoming, and enthusiastic about food and service. Use appropriate emojis for food and dining."
    elif industry.lower() in ['technology', 'software', 'it', 'startup']:
        industry_context = "IMPORTANT: This is a technology context. Be knowledgeable, innovative, and solution-focused. Use technical terms appropriately."
    elif industry.lower() in ['education', 'school', 'university', 'training']:
        industry_context = "IMPORTANT: This is an educational context. Be helpful, encouraging, and clear in explanations. Focus on learning and growth."
    
    # Personality-specific instructions
    personality_instructions = {
        'Professional': f"""• Maintain a formal, business-like tone
• Use clear, concise language
• Avoid slang or casual expressions
• Be respectful and courteous
• Focus on facts and solutions
{industry_context}""",
        
        'Friendly': f"""• Use a warm, approachable tone
• Be conversational and personable
• Use phrases like "I'd be happy to help" and "Let me assist you with that"
• Show genuine interest in helping
• Use exclamation points sparingly but appropriately
{industry_context}""",
        
        'Humorous': f"""• Use light, appropriate humor when suitable
• Include relevant emojis to enhance communication 😊
• Make jokes that are industry-appropriate and professional
• Use wordplay and friendly banter
• Keep humor positive and never offensive
• For healthcare: Use gentle, caring humor only
• For entertainment: Feel free to be more playful and fun
{industry_context}""",
        
        'Expert': f"""• Demonstrate deep knowledge and expertise
• Use technical terms when appropriate but explain them
• Provide detailed, comprehensive answers
• Show confidence in your knowledge
• Reference specific details and facts
• Be authoritative but not condescending
{industry_context}""",
        
        'Caring': f"""• Show empathy and understanding
• Use phrases like "I understand your concern" and "I'm here to help"
• Be patient and supportive
• Acknowledge feelings and concerns
• Use gentle, reassuring language
• Show genuine care for the customer's wellbeing
{industry_context}""",
        
        'Enthusiastic': f"""• Use energetic, positive language
• Show excitement about helping and the company
• Use exclamation points and positive expressions
• Be motivating and encouraging
• Use phrases like "That's fantastic!" and "I'm excited to help!"
• Show passion for the company and its services
{industry_context}""",
        
        'Formal': f"""• Use very formal, traditional language
• Avoid contractions (use "I will" instead of "I'll")
• Be extremely polite and respectful
• Use formal greetings and closings
• Maintain a serious, professional demeanor
• Use proper titles and formal address
{industry_context}""",
        
        'Casual': f"""• Use relaxed, informal language
• Use contractions naturally ("I'll", "you're", "we've")
• Be conversational and easy-going
• Use casual greetings and expressions
• Keep the tone light and approachable
• Use everyday language that's easy to understand
{industry_context}"""
    }
    
    return personality_instructions.get(tone, personality_instructions['Professional'])

def generate_dynamic_prompt(username, user_input, knowledge_context, max_words=20):
    """Generate a dynamic prompt based on company configuration"""
    try:
        company_config = get_company_config(username)
        
        if not company_config:
            # Fallback to basic prompt if no company config
            return f"""You are a helpful customer service assistant.
Answer the user's question briefly (max 6 lines): {user_input}

Business Information (use only if relevant; do not invent facts):
{knowledge_context}

Rules:
- Use ONLY the business information provided above to answer questions.
- If the answer is not in the business information above, say you don't have that specific information and ask a clarifying question.
- Keep the tone helpful and professional.
- Be specific and accurate based on the available business information only.
"""

        # Extract company information
        company_name = company_config.get('companyName', 'our company')
        company_url = company_config.get('companyUrl', '')
        company_description = company_config.get('companyDescription', '')
        tone = company_config.get('tone', 'Professional')
        industry = company_config.get('industry', '')
        response_length = company_config.get('responseLength', 50)
        
        # Get personality-specific instructions
        personality_instructions = get_personality_instructions(tone, industry)
        
        # Generate dynamic prompt based on company data
        prompt = f"""You are the official AI assistant for {company_name}.
{f"Company Website: {company_url}" if company_url else ""}
{f"Industry: {industry}" if industry else ""}
{f"About {company_name}: {company_description}" if company_description else ""}

Your role is to help customers with questions about {company_name} and provide accurate information based on the knowledge base provided.

PERSONALITY & COMMUNICATION STYLE:
{personality_instructions}

RESPONSE LENGTH: Keep your response to EXACTLY {max_words} words. Be concise but informative. Do not exceed this word count.

Answer the user's question: {user_input}

Knowledge Base Information (use only if relevant; do not invent facts):
{knowledge_context}

Rules:
- Always represent {company_name} professionally
- Use ONLY the business information provided above to answer questions
- If the answer is not in the knowledge base, politely say you don't have that specific information and offer to help with something else
- Follow the personality and communication style guidelines above
- Focus on helping customers with {company_name}'s services and information
- Be specific and accurate based on the available business information only
- If asked about {company_name}'s website, direct them to {company_url if company_url else "our official website"}
"""

        return prompt
    except Exception as e:
        logger.error(f"Error generating dynamic prompt: {e}")
        # Fallback to basic prompt
        return f"""You are a helpful customer service assistant.
Answer the user's question briefly (max 6 lines): {user_input}

Business Information (use only if relevant; do not invent facts):
{knowledge_context}

Rules:
- Use ONLY the business information provided above to answer questions.
- If the answer is not in the business information above, say you don't have that specific information and ask a clarifying question.
- Keep the tone helpful and professional.
- Be specific and accurate based on the available business information only.
"""

def get_chatgpt_response(user_input):
    try:
        # Check for common questions first
        user_input_lower = user_input.lower().strip()
        for question, response in COMMON_QUESTIONS.items():
            if question in user_input_lower:
                return response

        # Check cache
        cached_response = get_cached_response(user_input)
        if cached_response:
            return cached_response

        # Create a concise prompt using per-user knowledge base data
        knowledge_context = ""
        username = request.json.get('username') or ''
        kb = KB_CACHE.get(username) or load_knowledge_base_for(username)
        
        if kb and len(str(kb).strip()) > 0:
            # Convert knowledge base to a readable format regardless of structure
            knowledge_context = format_knowledge_base(kb)
            # Debug: Log the knowledge context for troubleshooting
            logger.info(f"Knowledge context for user {username}: {knowledge_context[:500]}...")
        else:
            logger.warning(f"No knowledge base found for user: {username}")
            # Try to load from any available user as fallback
            for cached_user, cached_data in KB_CACHE.items():
                if cached_data and len(str(cached_data).strip()) > 0:
                    logger.info(f"Using fallback knowledge base from user: {cached_user}")
                    knowledge_context = format_knowledge_base(cached_data)
                    break
        
        # Get response length from company config first
        company_config = get_company_config(username)
        max_words = company_config.get('responseLength', 20) if company_config else 20
        
        # Generate dynamic prompt based on company configuration with word count
        prompt = generate_dynamic_prompt(username, user_input, knowledge_context, max_words)

        response = model.generate_content(prompt)
        reply = response.text.strip()
        reply = reply.replace('*', '')
        
        # Enforce word count target: refine if too short, trim if too long
        words = reply.split()
        if len(words) < max_words - 2:
            try:
                refine_prompt = f"Rewrite the following assistant reply to EXACTLY {max_words} words without adding any new facts. Keep the same meaning and style. Reply only with the rewritten text.\n\nOriginal reply:\n{reply}"
                refined = model.generate_content(refine_prompt)
                reply = (refined.text or '').strip() or reply
            except Exception as _:
                pass
            words = reply.split()
        if len(words) > max_words:
            reply = ' '.join(words[:max_words])
            if not reply.endswith(('.', '!', '?')):
                reply += '.'
        
        # Limit to 6 lines
        lines = reply.splitlines()
        reply = '\n'.join(lines[:6])

        # Cache the response
        cache_response(user_input, reply)
        
        return reply
    except Exception as e:
        logger.error(f"Error calling Gemini API: {str(e)}")
        return "I apologize for the inconvenience, but I'm currently experiencing some technical difficulties. Please try again in a moment."

def handle_appointment_title(user_message):
    """Handle the first step of appointment scheduling - getting the title"""
    try:
        # Store the title and ask for date/time
        appointment_state['title'] = user_message.strip()
        appointment_state['waiting_for_title'] = False
        
        return f"Great! Your appointment title is: {appointment_state['title']}\n\nNow, please select the date and time for your appointment. I'll show you a calendar to choose from."
        
    except Exception as e:
        logger.error(f"Error handling appointment title: {str(e)}")
        return "Sorry, I encountered an error. Please try again."

def handle_appointment_scheduling(user_message):
    try:
        if not rtdb_available:
            return "Appointment storage is not configured. Please contact support."
        
        # Parse appointment details from user message
        lines = user_message.split('\n')
        title = None
        date_time = None
        
        for line in lines:
            line = line.strip()
            if line.lower().startswith('title:'):
                title = line.split(':', 1)[1].strip()
            elif line.lower().startswith('date') or line.lower().startswith('time'):
                date_time = line.split(':', 1)[1].strip()
        
        if not title or not date_time:
            return "Please provide both Title and Date/Time in the correct format:\nTitle: [Appointment Title]\nDate and Time: [Date and Time]"
        
        # Try to parse the date and time
        try:
            # Simple date parsing - you might want to use a more robust library like dateutil
            from datetime import datetime
            import re
            
            # Try to extract date and time from the string
            # This is a simple parser - you might want to enhance it
            appointment_time = datetime.now()  # Default to now, you can enhance this parsing
            
            # Generate a unique ID
            appointment_id = f"APT-{int(time.time())}-{uuid.uuid4().hex[:8]}"
            
            # Create appointment object
            appointment = {
                'id': appointment_id,
                'title': title,
                'time': appointment_time.isoformat(),
                'status': 'pending',
                'created_at': int(time.time() * 1000)
            }
            
            # Save to Firebase Realtime Database
            fb_db.reference('appointments').child(appointment_id).set(appointment)
            
            return f"Great! I've scheduled your appointment:\n\nTitle: {title}\nDate/Time: {date_time}\nAppointment ID: {appointment_id}\n\nPlease save this appointment ID for future reference."
            
        except Exception as e:
            logger.error(f"Error parsing appointment time: {str(e)}")
            return "I had trouble parsing the date and time. Please use a clear format like 'January 15, 2024 at 2:00 PM'."
            
    except Exception as e:
        logger.error(f"Error handling appointment scheduling: {str(e)}")
        return "Sorry, I encountered an error while scheduling your appointment. Please try again."

def handle_appointment_cancellation(user_message, username):
    try:
        if not rtdb_available:
            return "Appointment storage is not configured. Please contact support."
        
        # Extract appointment ID from user message
        appointment_id = None
        
        # Look for APT- pattern
        import re
        apt_match = re.search(r'APT-[A-Za-z0-9-]+', user_message)
        if apt_match:
            appointment_id = apt_match.group()
        else:
            # Look for "appointment id:" pattern
            lines = user_message.split('\n')
            for line in lines:
                line = line.strip()
                if line.lower().startswith('appointment id:'):
                    appointment_id = line.split(':', 1)[1].strip()
                    break
        
        if not appointment_id:
            return "Please provide your appointment ID. It should look like 'APT-1234567890-abc12345'."
        
        # Check if appointment exists and update status (scoped to user and bot if available)
        try:
            bot_id = (request.json.get('bot_id') if request.is_json else (request.args.get('botId') or '')) or ''
            base = _base_path(username, bot_id)
            appointment_ref = fb_db.reference(f'{base}/appointments').child(appointment_id)
            appointment = appointment_ref.get()
            
            if not appointment:
                return f"Appointment ID {appointment_id} not found. Please check your appointment ID and try again."
            
            # Update appointment status
            appointment_ref.update({'status': 'cancelled'})
            
            # Also free the slot lock if present
            try:
                appt_time = appointment.get('time') if isinstance(appointment, dict) else None
                if appt_time:
                    appt_dt = datetime.fromisoformat(appt_time.replace('Z', '+00:00')).astimezone(timezone.utc)
                    slot_key = appt_dt.strftime('%Y%m%d-%H%M')
                    fb_db.reference(f'{base}/slot_locks').child(slot_key).set({'status': 'cancelled', 'appointment_id': appointment_id})
            except Exception as _e:
                logger.warning(f"Failed to free slot lock: {_e}")

            return f"Appointment {appointment_id} has been cancelled successfully."
            
        except Exception as e:
            logger.error(f"Error cancelling appointment: {str(e)}")
            return "Sorry, I encountered an error while cancelling your appointment. Please try again."
            
    except Exception as e:
        logger.error(f"Error handling appointment cancellation: {str(e)}")
        return "Sorry, I encountered an error while processing your cancellation request. Please try again."

@app.route('/')
def landing_page():
    # Serve the marketing landing page
    return render_template('landing.html')

@app.route('/login')
def login_page():
    # Serve the login page
    return render_template('login.html')

@app.route('/signup')
def signup_page():
    # Serve the signup page
    return render_template('signup.html')

# Serve the robot image from project root so we can use it on landing page
@app.route('/assets/robot.png')
def robot_asset():
    try:
        root_dir = os.path.dirname(__file__)
        return send_from_directory(root_dir, 'chatbot-robot-isolated-white-background.png')
    except Exception as _e:
        logger.warning(f"Failed to serve robot image: {_e}")
        # Fallback to 404
        return ('', 404)

@app.route('/assets/login-robot.png')
def login_robot_asset():
    try:
        root_dir = os.path.dirname(__file__)
        return send_from_directory(root_dir, 'futuristic-robot-illustration.png')
    except Exception as _e:
        logger.warning(f"Failed to serve login robot image: {_e}")
        return ('', 404)

@app.route('/assets/avatar-smile.png')
def avatar_smile_asset():
    try:
        root_dir = os.path.dirname(__file__)
        return send_from_directory(root_dir, 'smiling-young-man-illustration.png')
    except Exception as _e:
        logger.warning(f"Failed to serve avatar image: {_e}")
        return ('', 404)

@app.route('/assets/avatar-cowboy.png')
def avatar_cowboy_asset():
    try:
        root_dir = os.path.dirname(__file__)
        return send_from_directory(root_dir, 'happy-smiley-boy-farmer-cowboy-hat.png')
    except Exception as _e:
        logger.warning(f"Failed to serve cowboy avatar image: {_e}")
        return ('', 404)

@app.route('/assets/avatar-doctor.png')
def avatar_doctor_asset():
    try:
        root_dir = os.path.dirname(__file__)
        return send_from_directory(root_dir, 'portrait-3d-female-doctor.png')
    except Exception as _e:
        logger.warning(f"Failed to serve doctor avatar image: {_e}")
        return ('', 404)

@app.route('/knowledgebase')
def setup_knowledge_page():
    # Check if user has token in localStorage by checking query params or headers
    token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    
    if not token:
        # If no token provided, redirect to login
        return render_template('login.html')
    
    # Verify token
    username = _verify_token(token)
    if not username:
        return render_template('login.html')
    
    return render_template('setup_knowledge.html')

@app.route('/scripttag/<username>')
def setup_customize_page(username):
    # Require token and ensure it matches the route username
    token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    verified = _verify_token(token) if token else ''
    if not verified or verified != (username or '').strip():
        return render_template('login.html')
    return render_template('setup_customize.html')

# Token-based clean routes that avoid exposing username in the URL
@app.route('/scripttag')
def setup_customize_page_token():
    token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    username = _verify_token(token) if token else ''
    if not username:
        return render_template('login.html')
    return setup_customize_page(username)

# Removed test/demo routes and sample pages for production cleanup

@app.route('/generate_script/<username>')
def generate_script(username):
    """Generate the correct script tag for a user's website"""
    try:
        # Require token and ensure subject matches
        token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
        subject = _verify_token(token) if token else ''
        if not subject or subject != (username or '').strip():
            return jsonify({'error': 'Unauthorized'}), 401
        # Get company configuration
        config = get_company_config(username)
        if not config:
            return jsonify({'error': 'Company configuration not found for user: ' + username}), 404
        
        # Generate the script tag with proper configuration
        script_tag = f'''<!-- Guru Restaurants Chatbot Widget -->
<script>
    // Set the username for the chatbot
    window.CHATBOT_USERNAME = '{username}';
    
    // Get authentication token and set it in localStorage
    async function setupAuth() {{
        try {{
            const response = await fetch('http://localhost:5001/auth/login', {{
                method: 'POST',
                headers: {{ 'Content-Type': 'application/json' }},
                body: JSON.stringify({{ username: '{username}', password: 'password123' }})
            }});
            const data = await response.json();
            if (data.success && data.token) {{
                localStorage.setItem('token', data.token);
                localStorage.setItem('username', data.username);
                console.log('✅ Authentication successful!');
            }} else {{
                console.error('❌ Authentication failed:', data);
            }}
        }} catch (error) {{
            console.error('❌ Authentication error:', error);
        }}
    }}
    
    // Setup auth when page loads
    setupAuth();
</script>
<script src="http://localhost:5001/static/chatbot-widget.js"></script>
<script>
    // Update widget configuration with company settings
    ChatbotWidget.updateConfig({{
        apiUrl: 'http://localhost:5001',
        username: '{username}',
        primaryColor: '{config.get('primaryColor', '#4f46e5')}',
        secondaryColor: '#7c3aed'
    }});
</script>'''
        
        return script_tag, 200, {'Content-Type': 'text/html; charset=utf-8'}
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate_script')
def generate_script_token():
    token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    username = _verify_token(token) if token else ''
    if not username:
        return jsonify({'error': 'Unauthorized'}), 401
    return generate_script(username)

@app.route('/dashboard/<username>')
def user_dashboard_page(username):
    # Require token and ensure it matches the route username
    token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    verified_username = _verify_token(token) if token else ''
    if not verified_username or verified_username != username:
        return render_template('login.html')
    try:
        username = (username or '').strip()
        if not username:
            return render_template('user_dashboard.html', username='', leads=[], appointments=[], conversations=[], company_name='IM Solutions')

        leads = []
        appointments = []
        conversations = []
        bot_id = None
        
        if rtdb_available:
            # Get user's single bot ID
            try:
                bots_ref = fb_db.reference(f'{username}/bots')
                bots = bots_ref.get() or {}
                if bots and len(bots) > 0:
                    bot_id = list(bots.keys())[0]  # Get the first (and only) bot ID
            except Exception as e:
                logger.warning(f"Bot fetch failed: {e}")
            
            # Try to load data from new structure first (with bot_id), then fallback to old structure
            if bot_id:
                # Try new structure: {username}/bots/{bot_id}/leads
                try:
                    base = _base_path(username, bot_id)
                    lref = fb_db.reference(f'{base}/leads')
                    lraw = lref.get() or {}
                    if isinstance(lraw, dict):
                        leads = list(lraw.values())
                except Exception as e:
                    logger.warning(f"New structure leads fetch failed: {e}")
                
                try:
                    base = _base_path(username, bot_id)
                    aref = fb_db.reference(f'{base}/appointments')
                    araw = aref.get() or {}
                    if isinstance(araw, dict):
                        appointments = list(araw.values())
                except Exception as e:
                    logger.warning(f"New structure appointments fetch failed: {e}")
                
                try:
                    base = _base_path(username, bot_id)
                    cref = fb_db.reference(f'{base}/conversations')
                    craw = cref.get() or {}
                    if isinstance(craw, dict):
                        conversations = list(craw.values())
                except Exception as e:
                    logger.warning(f"New structure conversations fetch failed: {e}")
            
            # If no data found in new structure, try old structure: {username}/leads
            if not leads:
                try:
                    lref = fb_db.reference(f'{username}/leads')
                    lraw = lref.get() or {}
                    if isinstance(lraw, dict):
                        leads = list(lraw.values())
                        logger.info(f"Loaded {len(leads)} leads from old structure")
                except Exception as e:
                    logger.warning(f"Old structure leads fetch failed: {e}")
            
            if not appointments:
                try:
                    aref = fb_db.reference(f'{username}/appointments')
                    araw = aref.get() or {}
                    if isinstance(araw, dict):
                        appointments = list(araw.values())
                        logger.info(f"Loaded {len(appointments)} appointments from old structure")
                except Exception as e:
                    logger.warning(f"Old structure appointments fetch failed: {e}")
            
            if not conversations:
                try:
                    cref = fb_db.reference(f'{username}/conversations')
                    craw = cref.get() or {}
                    if isinstance(craw, dict):
                        conversations = list(craw.values())
                        logger.info(f"Loaded {len(conversations)} conversations from old structure")
                except Exception as e:
                    logger.warning(f"Old structure conversations fetch failed: {e}")
        # Determine role for this username
        token = request.args.get('token') or ''
        is_subadmin = False
        parent_username = ''
        try:
            if rtdb_available:
                uref = fb_db.reference('users').child(username)
                uobj = uref.get() or {}
                is_subadmin = (isinstance(uobj, dict) and uobj.get('role') == 'subadmin')
                parent_username = (uobj.get('parent') or '') if isinstance(uobj, dict) else ''
        except Exception:
            is_subadmin = False
            parent_username = ''

        # Enforce onboarding: if main user has no company_config, redirect to setup first
        try:
            if not is_subadmin:
                cfg = get_company_config(username)
                if not cfg:
                    return redirect("/setup")
        except Exception:
            pass

        # If sub-admin, load data from parent account and filter assigned leads
        if is_subadmin and parent_username:
            try:
                # Reset lists and pull from parent context
                leads = []
                appointments = []
                conversations = []
                bot_id = None
                if rtdb_available:
                    # Backfill any unassigned leads first
                    try:
                        backfill_unassigned_leads(parent_username)
                    except Exception as _e:
                        logger.warning(f"Backfill failed: {_e}")
                    try:
                        bots_ref = fb_db.reference(f'{parent_username}/bots')
                        bots = bots_ref.get() or {}
                        if bots and len(bots) > 0:
                            bot_id = list(bots.keys())[0]
                    except Exception as _e:
                        logger.warning(f"Bot fetch failed (subadmin parent): {_e}")
                    # Backfill any unassigned appointments as well
                    try:
                        backfill_unassigned_appointments(parent_username)
                    except Exception as _e:
                        logger.warning(f"Backfill appointments failed: {_e}")
                    if bot_id:
                        try:
                            base = _base_path(parent_username, bot_id)
                            lref = fb_db.reference(f'{base}/leads')
                            lraw = lref.get() or {}
                            if isinstance(lraw, dict):
                                leads = list(lraw.values())
                        except Exception as _e:
                            logger.warning(f"Parent leads fetch failed: {_e}")
                        try:
                            base = _base_path(parent_username, bot_id)
                            aref = fb_db.reference(f'{base}/appointments')
                            araw = aref.get() or {}
                            if isinstance(araw, dict):
                                appointments = list(araw.values())
                        except Exception as _e:
                            logger.warning(f"Parent appointments fetch failed: {_e}")
                        try:
                            base = _base_path(parent_username, bot_id)
                            cref = fb_db.reference(f'{base}/conversations')
                            craw = cref.get() or {}
                            if isinstance(craw, dict):
                                conversations = list(craw.values())
                        except Exception as _e:
                            logger.warning(f"Parent conversations fetch failed: {_e}")
                    # Fallbacks to old structure
                    if not leads:
                        try:
                            lref = fb_db.reference(f'{parent_username}/leads')
                            lraw = lref.get() or {}
                            if isinstance(lraw, dict):
                                leads = list(lraw.values())
                        except Exception as _e:
                            logger.warning(f"Parent old leads fetch failed: {_e}")
                    if not appointments:
                        try:
                            aref = fb_db.reference(f'{parent_username}/appointments')
                            araw = aref.get() or {}
                            if isinstance(araw, dict):
                                appointments = list(araw.values())
                        except Exception as _e:
                            logger.warning(f"Parent old appointments fetch failed: {_e}")
                    if not conversations:
                        try:
                            cref = fb_db.reference(f'{parent_username}/conversations')
                            craw = cref.get() or {}
                            if isinstance(craw, dict):
                                conversations = list(craw.values())
                        except Exception as _e:
                            logger.warning(f"Parent old conversations fetch failed: {_e}")
                # Only keep leads/appointments assigned to this sub-admin
                leads = [l for l in leads if isinstance(l, dict) and (l.get('assigned_to') == username)]
                appointments = [a for a in appointments if isinstance(a, dict) and (a.get('assigned_to') == username)]
                
                # Filter conversations to only show those from assigned leads
                assigned_session_ids = {l.get('session_id') for l in leads if l.get('session_id')}
                conversations = [c for c in conversations if isinstance(c, dict) and c.get('session_id') in assigned_session_ids]
            except Exception as _e:
                logger.warning(f"Failed sub-admin data load: {_e}")

        # sort newest first by timestamp/created_at
        leads.sort(key=lambda x: x.get('created_at', 0), reverse=True)
        appointments.sort(key=lambda x: x.get('created_at', 0), reverse=True)
        conversations.sort(key=lambda x: x.get('timestamp', 0), reverse=True)

        # Get company configuration for the dashboard title
        company_config = get_company_config(username)
        company_name = company_config.get('companyName', 'IM Solutions') if company_config else 'IM Solutions'

        # Provide 'now' for overdue status comparisons in templates
        now_iso = datetime.now(timezone.utc).isoformat()
        return render_template('user_dashboard.html', username=username, leads=leads, appointments=appointments, conversations=conversations, token=token, is_subadmin=is_subadmin, company_name=company_name, now=now_iso)
    except Exception as e:
        logger.error(f"Dashboard error: {e}")
        return render_template('user_dashboard.html', username='', leads=[], appointments=[], conversations=[], company_name='IM Solutions', now=datetime.now(timezone.utc).isoformat())

@app.route('/feedback', methods=['POST'])
def save_feedback():
    try:
        data = request.get_json(force=True, silent=True) or {}
        username = data.get('username') or 'anonymous'
        rating = int(data.get('rating') or 0)
        reason = (data.get('reason') or '').strip()
        message = (data.get('message') or '').strip()
        session_id = data.get('session_id') or ''
        if rating < 1 or rating > 5:
            return jsonify({'success': False, 'error': 'invalid rating'}), 400
        # Persist to Firebase if available
        try:
            if rtdb_available:
                ref = fb_db.reference(f'{username}/feedback')
                item = { 'rating': rating, 'session_id': session_id, 'reason': reason, 'message': message, 'created_at': int(time.time()*1000) }
                ref.push(item)
        except Exception:
            pass

        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"feedback error: {e}")
        return jsonify({'success': False}), 500

@app.route('/dashboard')
def user_dashboard_page_token():
    token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    username = _verify_token(token) if token else ''
    if not username:
        return render_template('login.html')
    return user_dashboard_page(username)

@app.route('/feedback_summary/<username>', methods=['GET'])
def feedback_summary(username: str):
    """Return average rating and counts for customer satisfaction gauge.
    Accepts optional token via header/query/body but does not enforce strict auth for now.
    """
    try:
        avg = 0.0
        total = 0
        counts = [0,0,0,0,0]
        if rtdb_available:
            try:
                items = fb_db.reference(f'{username}/feedback').get() or {}
                if isinstance(items, dict):
                    for _, it in items.items():
                        try:
                            r = int((it.get('rating') or 0))
                        except Exception:
                            r = 0
                        if 1 <= r <= 5:
                            counts[r-1] += 1
                            total += 1
            except Exception:
                pass
        if total > 0:
            s = 0
            for i,c in enumerate(counts):
                s += (i+1) * c
            avg = s / total
        return jsonify({ 'success': True, 'average': avg, 'total': total, 'counts': counts })
    except Exception as e:
        logger.error(f"feedback_summary error: {e}")
        return jsonify({'success': False}), 500
@app.route('/hourly_activity/<username>')
def hourly_activity(username):
    """Return hourly conversation activity for a given day.

    Query params:
      - date: YYYY-MM-DD (required)
    Auth:
      - Authorization: Bearer <token> or token query param (optional if dashboard is public)
    Response JSON:
      { labels: ["12:00 AM", ...], messages: [..24..], users: [..24..] }
    """
    from datetime import datetime, timezone
    try:
        # Auth (best-effort, same as dashboard)
        token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
        if token:
            verified_username = _verify_token(token)
            if not verified_username or verified_username != username:
                return jsonify({'error': 'unauthorized'}), 401

        date_str = (request.args.get('date') or '').strip()
        if not date_str:
            return jsonify({'error': 'missing date'}), 400

        # Compute start/end of the requested UTC day
        try:
            # Treat incoming date as local calendar day but compute in UTC by assuming midnight local ~ easier use naive
            # Use datetime.fromisoformat to parse YYYY-MM-DD
            d = datetime.fromisoformat(date_str)
        except Exception:
            return jsonify({'error': 'invalid date'}), 400

        start_dt = datetime(d.year, d.month, d.day, 0, 0, 0, tzinfo=timezone.utc)
        end_dt = datetime(d.year, d.month, d.day, 23, 59, 59, 999000, tzinfo=timezone.utc)
        start_ms = int(start_dt.timestamp() * 1000)
        end_ms = int(end_dt.timestamp() * 1000)

        hourly_messages = [0] * 24
        hourly_user_sets = [set() for _ in range(24)]

        conversations = []
        bot_id = None
        if rtdb_available:
            try:
                # Find bot id
                bots_ref = fb_db.reference(f'{username}/bots')
                bots = bots_ref.get() or {}
                if bots and len(bots) > 0:
                    bot_id = list(bots.keys())[0]
            except Exception:
                bot_id = None

            # Try new structure
            try:
                if bot_id:
                    base = _base_path(username, bot_id)
                    cref = fb_db.reference(f'{base}/conversations')
                    craw = cref.get() or {}
                    if isinstance(craw, dict):
                        conversations = list(craw.values())
            except Exception:
                pass

            # Fallback to old structure
            if not conversations:
                try:
                    cref = fb_db.reference(f'{username}/conversations')
                    craw = cref.get() or {}
                    if isinstance(craw, dict):
                        conversations = list(craw.values())
                except Exception:
                    conversations = []

        # Aggregate by hour within the date window
        for c in conversations:
            try:
                ts = int(c.get('timestamp') or 0)
            except Exception:
                ts = 0
            if not ts:
                continue
            if ts < start_ms or ts > end_ms:
                continue
            try:
                dt = datetime.fromtimestamp(ts / 1000.0, tz=timezone.utc)
                hour = dt.hour
            except Exception:
                continue
            hourly_messages[hour] += 1
            sid = c.get('session_id') or ''
            if sid:
                hourly_user_sets[hour].add(sid)

        hourly_users = [len(s) for s in hourly_user_sets]
        labels = []
        for i in range(24):
            h = (i % 12) or 12
            period = 'AM' if i < 12 else 'PM'
            labels.append(f"{h}:00 {period}")

        return jsonify({'labels': labels, 'messages': hourly_messages, 'users': hourly_users})
    except Exception as e:
        logger.error(f"hourly_activity error: {e}")
        return jsonify({'error': 'internal error'}), 500
@app.route('/leads/<username>')
def leads_page(username):
    # Check if user has token
    token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    if not token:
        return render_template('login.html')
    verified_username = _verify_token(token)
    if not verified_username or verified_username != username:
        return render_template('login.html')
    try:
        leads = []
        bot_id = None
        
        if rtdb_available:
            # Get user's single bot ID
            try:
                bots_ref = fb_db.reference(f'{username}/bots')
                bots = bots_ref.get() or {}
                if bots and len(bots) > 0:
                    bot_id = list(bots.keys())[0]
            except Exception as e:
                logger.warning(f"Bot fetch failed: {e}")
            
            # Try new structure first, then fallback to old structure
            if bot_id:
                try:
                    base = _base_path(username, bot_id)
                    lref = fb_db.reference(f'{base}/leads')
                    lraw = lref.get() or {}
                    if isinstance(lraw, dict):
                        leads = list(lraw.values())
                except Exception as e:
                    logger.warning(f"New structure leads fetch failed: {e}")
            
            # Fallback to old structure
            if not leads:
                try:
                    lref = fb_db.reference(f'{username}/leads')
                    lraw = lref.get() or {}
                    if isinstance(lraw, dict):
                        leads = list(lraw.values())
                        logger.info(f"Loaded {len(leads)} leads from old structure")
                except Exception as e:
                    logger.error(f"Error loading leads for {username}: {e}")
        
        leads.sort(key=lambda x: x.get('created_at', 0), reverse=True)
        return render_template('leads.html', username=username, leads=leads)
    except Exception as e:
        logger.error(f"Leads page error: {e}")
        return render_template('leads.html', username=username, leads=[])

@app.route('/appointments/<username>')
def appointments_page(username):
    # Auth check
    token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    if not token:
        return render_template('login.html')
    verified_username = _verify_token(token)
    if not verified_username or verified_username != username:
        return render_template('login.html')
    try:
        appointments = []
        bot_id = None
        
        if rtdb_available:
            # Get user's single bot ID
            try:
                bots_ref = fb_db.reference(f'{username}/bots')
                bots = bots_ref.get() or {}
                if bots and len(bots) > 0:
                    bot_id = list(bots.keys())[0]
            except Exception as e:
                logger.warning(f"Bot fetch failed: {e}")
            
            # Try new structure first, then fallback to old structure
            if bot_id:
                try:
                    base = _base_path(username, bot_id)
                    aref = fb_db.reference(f'{base}/appointments')
                    araw = aref.get() or {}
                    if isinstance(araw, dict):
                        appointments = list(araw.values())
                except Exception as e:
                    logger.warning(f"New structure appointments fetch failed: {e}")
            
            # Fallback to old structure
            if not appointments:
                try:
                    aref = fb_db.reference(f'{username}/appointments')
                    araw = aref.get() or {}
                    if isinstance(araw, dict):
                        appointments = list(araw.values())
                        logger.info(f"Loaded {len(appointments)} appointments from old structure")
                except Exception as e:
                    logger.error(f"Error loading appointments for {username}: {e}")
        
        appointments.sort(key=lambda x: x.get('created_at', 0), reverse=True)
        return render_template('appointments.html', username=username, appointments=appointments, token=token)
    except Exception as e:
        logger.error(f"Appointments page error: {e}")
        return render_template('appointments.html', username=username, appointments=[], token=token)

# -----------------------------
# Jinja filters and utilities
# -----------------------------
@app.template_filter('format_time')
def format_time_filter(value, fmt='%d %b %Y, %I:%M %p'):
    """Format ISO8601 time strings stored in RTDB to a readable local time.
    Accepts values like '2025-09-17T14:30:00+00:00' or '2025-09-17T14:30:00Z'.
    """
    if not value:
        return ''
    try:
        s = str(value)
        if s.endswith('Z'):
            dt = datetime.fromisoformat(s.replace('Z', '+00:00'))
        else:
            dt = datetime.fromisoformat(s)
        # Convert to local time for display
        local_dt = dt.astimezone()
        return local_dt.strftime(fmt)
    except Exception:
        # Fallback: return as-is
        return str(value)

@app.route('/bots/<username>')
def bots_page(username):
    # Auth check
    token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    if not token:
        return render_template('login.html')
    verified_username = _verify_token(token)
    if not verified_username or verified_username != username:
        return render_template('login.html')
    try:
        bots = []
        if rtdb_available:
            bref = fb_db.reference(f'{username}/bots')
            braw = bref.get() or {}
            if isinstance(braw, dict):
                bots = list(braw.values())
        # newest first
        bots.sort(key=lambda x: x.get('created_at', 0), reverse=True)
        return render_template('bots.html', username=username, bots=bots, token=token)
    except Exception as e:
        logger.error(f"Bots page error: {e}")
        return render_template('bots.html', username=username, bots=[], token=token)

@app.route('/bots/create', methods=['POST'])
@require_auth
def bots_create():
    try:
        if not rtdb_available:
            return jsonify({'success': False, 'message': 'Database not configured.'}), 503
        data = request.json or {}
        username = (data.get('username') or '').strip() or request.current_user
        name = (data.get('name') or 'My Chatbot').strip()
        
        # Check if user already has a bot
        existing_bots_ref = fb_db.reference(f'{username}/bots')
        existing_bots = existing_bots_ref.get() or {}
        
        if existing_bots and len(existing_bots) > 0:
            # User already has a bot, return the existing one
            existing_bot = list(existing_bots.values())[0]
            return jsonify({'success': True, 'bot': existing_bot, 'message': 'You already have a chatbot. Redirecting to setup...'})
        
        # Create new bot
        bot_id = str(uuid.uuid4())
        bot = {
            'id': bot_id,
            'name': name,
            'created_at': int(time.time() * 1000)
        }
        fb_db.reference(f'{username}/bots').child(bot_id).set(bot)
        return jsonify({'success': True, 'bot': bot})
    except Exception as e:
        logger.error(f"Create bot error: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/bots/update', methods=['POST'])
@require_auth
def bots_update():
    try:
        if not rtdb_available:
            return jsonify({'success': False, 'message': 'Database not configured.'}), 503
        data = request.json or {}
        username = (data.get('username') or '').strip() or request.current_user
        bot_id = (data.get('id') or '').strip()
        name = (data.get('name') or '').strip()
        if not bot_id or not name:
            return jsonify({'success': False, 'message': 'Missing id or name'}), 400
        ref = fb_db.reference(f'{username}/bots').child(bot_id)
        if not ref.get():
            return jsonify({'success': False, 'message': 'Bot not found'}), 404
        ref.update({'name': name})
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Update bot error: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/bots/delete', methods=['POST'])
@require_auth
def bots_delete():
    try:
        if not rtdb_available:
            return jsonify({'success': False, 'message': 'Database not configured.'}), 503
        data = request.json or {}
        username = (data.get('username') or '').strip() or request.current_user
        bot_id = (data.get('id') or '').strip()
        if not bot_id:
            return jsonify({'success': False, 'message': 'Missing id'}), 400
        
        # Check if bot exists
        ref = fb_db.reference(f'{username}/bots').child(bot_id)
        if not ref.get():
            return jsonify({'success': False, 'message': 'Bot not found'}), 404
        
        # Delete all associated data
        try:
            # Delete from new structure: {username}/bots/{bot_id}/...
            base = _base_path(username, bot_id)
            fb_db.reference(f'{base}/leads').delete()
            fb_db.reference(f'{base}/appointments').delete()
            fb_db.reference(f'{base}/conversations').delete()
            logger.info(f"Deleted data from new structure for bot {bot_id}")
        except Exception as e:
            logger.warning(f"Failed to delete from new structure: {e}")
        
        # Also try to delete from old structure as fallback
        try:
            fb_db.reference(f'{username}/leads').delete()
            fb_db.reference(f'{username}/appointments').delete()
            fb_db.reference(f'{username}/conversations').delete()
            logger.info(f"Deleted data from old structure for user {username}")
        except Exception as e:
            logger.warning(f"Failed to delete from old structure: {e}")
        
        # Delete the bot itself
        ref.delete()
        
        # Also delete company config and knowledge base
        try:
            fb_db.reference(f'{username}/company_config').delete()
            fb_db.reference(f'{username}/knowledge_base').delete()
            logger.info(f"Deleted company config and knowledge base for user {username}")
        except Exception as e:
            logger.warning(f"Failed to delete company config/knowledge base: {e}")
        
        return jsonify({'success': True, 'message': 'Chatbot and all associated data deleted successfully'})
    except Exception as e:
        logger.error(f"Delete bot error: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500

def has_user_provided_contact_info(username, session_id):
    """Check if user has already provided contact information in this session"""
    try:
        session_key = f"{username}_{session_id}"
        # If this session already finished lead capture, skip asking again
        if lead_capture_completed.get(session_key, False):
            logger.info(f"Contact info check for {username}: lead capture already completed in this session, returning True")
            return True
        # If we're mid-capture, we haven't collected yet
        if session_key in lead_capture_state:
            logger.info(f"Contact info check for {username}: user is in lead capture mode, returning False")
            return False
        # For a new session, always allow prompting (do not suppress based on prior leads)
        logger.info(f"Allowing lead capture for new session: {session_key}")
        return False
        
    except Exception as e:
        logger.error(f"Error checking contact info: {e}")
        return False

def get_message_count_simple(username, session_id):
    """Get the number of messages sent by user in this session using in-memory counter"""
    try:
        session_key = f"{username}_{session_id}"
        return message_counters.get(session_key, 0)
    except Exception as e:
        logger.error(f"Error getting message count: {e}")
        return 0

def increment_message_count(username, session_id):
    """Increment the message count for this user session"""
    try:
        session_key = f"{username}_{session_id}"
        message_counters[session_key] = message_counters.get(session_key, 0) + 1
        logger.info(f"Message count for {session_key}: {message_counters[session_key]}")
    except Exception as e:
        logger.error(f"Error incrementing message count: {e}")

def get_message_count(username, session_id):
    """Get the number of messages sent by user in this session"""
    try:
        if not rtdb_available:
            return 0
        
        # Get conversations for this user and session
        conversations_ref = fb_db.reference(f'{username}/conversations')
        conversations = conversations_ref.get() or {}
        
        # Count messages in this session
        count = 0
        for conv_id, conv_data in conversations.items():
            if conv_data.get('session_id') == session_id:
                count += 1
        
        return count
    except Exception as e:
        logger.error(f"Error getting message count: {e}")
        return 0

def is_lead_capture_message(user_message):
    """Check if the user message is providing lead information"""
    msg_lower = user_message.lower()
    
    # Check for lead information patterns
    lead_patterns = [
        'name:', 'email:', 'phone:', 'contact:',
        'my name is', 'my email is', 'my phone is',
        'call me', 'reach me at', 'contact me at',
        '@', '.com', '.org', '.net',  # email patterns
        '+1', '+91', '+44',  # phone patterns
        'name is', 'email is', 'phone is'
    ]
    
    return any(pattern in msg_lower for pattern in lead_patterns)

def generate_lead_capture_message(username, step=1):
    """Generate a lead capture message based on company configuration and personality"""
    try:
        company_config = get_company_config(username)
        if not company_config:
            return None
        
        tone = company_config.get('tone', 'Professional')
        industry = company_config.get('industry', '')
        company_name = company_config.get('companyName', 'our company')
        
        # Step 1: Ask for name (one or two words)
        if step == 1:
            name_messages = {
                'Professional': f"To personalize your experience, please share your name.",
                'Friendly': f"Awesome! Could you share your name? 😊",
                'Humorous': f"Roll call 😄 What's your name?",
                'Expert': f"For accurate records with {company_name}, please provide your name.",
                'Caring': f"I'd love to address you properly 🤗 Please share your name.",
                'Enthusiastic': f"Great! 🚀 What's your name?",
                'Formal': f"May I kindly request your name for our records?",
                'Casual': f"Hey! What's your name?"
            }
            return name_messages.get(tone, name_messages['Professional'])
        
        # Step 2: Ask for phone
        elif step == 2:
            phone_messages = {
                'Professional': f"Thank you! Could you also share your phone number so I can reach you if needed?",
                'Friendly': f"Great! 😊 What's your phone number? I'd love to stay in touch!",
                'Humorous': f"Perfect! 😄 What's your phone number? Don't worry, I won't call you at 3 AM... probably!",
                'Expert': f"Excellent. I'd also appreciate your phone number for important updates about {company_name}.",
                'Caring': f"Wonderful! 🤗 What's your phone number? I want to make sure I can reach you when it matters.",
                'Enthusiastic': f"Awesome! 🚀 What's your phone number? I'm so excited to have your contact details!",
                'Formal': f"Thank you. May I also request your telephone number for important communications?",
                'Casual': f"Cool! What's your phone number? I'll make sure to keep you updated!"
            }
            # Allow skip
            if str(request.json.get('message','')).strip().lower() == 'skip':
                return phone_messages.get(tone, phone_messages['Professional'])
            return phone_messages.get(tone, phone_messages['Professional'])
        
        # Step 3: Ask for email
        elif step == 3:
            email_messages = {
                'Professional': f"Perfect! One last thing - could you share your email address so I can send you updates?",
                'Friendly': f"Great! 😊 Last question - what's your email address? I'd love to keep you updated!",
                'Humorous': f"Almost there! 😄 What's your email address? I promise not to spam you - only the good stuff!",
                'Expert': f"Excellent. Finally, I'd appreciate your email address for detailed updates about {company_name}.",
                'Caring': f"Wonderful! 🤗 One last thing - what's your email address? I want to make sure you get all the important updates.",
                'Enthusiastic': f"Fantastic! 🚀 Last step - what's your email address? I can't wait to send you exciting updates!",
                'Formal': f"Thank you. Finally, may I request your email address for relevant communications?",
                'Casual': f"Awesome! Last thing - what's your email address? I'll make sure to keep you in the loop!"
            }
            return email_messages.get(tone, email_messages['Professional'])
        
        return None
        
    except Exception as e:
        logger.error(f"Error generating lead capture message: {e}")
        return None

def handle_lead_capture_step(user_message, username, session_key, bot_id=''):
    """Handle step-by-step lead capture process"""
    try:
        import re
        
        state = lead_capture_state[session_key]
        current_step = state['step']
        
        # Extract information based on current step
        if current_step == 1:  # Name (allow 1 or 2 words; trim if more)
            name_patterns = [
                r'name[:\s]+([a-zA-Z\s]+)',
                r'my name is ([a-zA-Z\s]+)',
                r'i\'m ([a-zA-Z\s]+)',
                r'call me ([a-zA-Z\s]+)',
                r'^([a-zA-Z\s]+)$'  # Just the name itself
            ]
            
            name = None
            for pattern in name_patterns:
                match = re.search(pattern, user_message, re.IGNORECASE)
                if match:
                    name = match.group(1).strip()
                    break
            
            # If no pattern matched, check if it's just a simple name (letters and spaces)
            if not name and re.match(r'^[a-zA-Z\s]+$', user_message.strip()) and len(user_message.strip()) > 1:
                name = user_message.strip()
            
            # Enforce max two words (allow one or two)
            if name:
                parts = [p for p in name.split() if p]
                if len(parts) >= 3:
                    name = parts[0] + ' ' + parts[1]
            
            if name and len(name) > 1:
                state['name'] = name
                state['step'] = 2
                return generate_lead_capture_message(username, 2)
            else:
                return "I didn't catch your full name. Please provide first and last name."
        
        elif current_step == 2:  # Phone
            phone_patterns = [
                r'phone[:\s]+([+\d\s\-\(\)]+)',
                r'call[:\s]+([+\d\s\-\(\)]+)',
                r'(\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})'
            ]
            
            phone = None
            # Allow skip at this step
            if user_message.strip().lower() == 'skip':
                state['phone'] = ''
                state['step'] = 3
                return generate_lead_capture_message(username, 3)
            for pattern in phone_patterns:
                match = re.search(pattern, user_message)
                if match:
                    candidate = match.group(1).strip()
                    digits_only_match = re.sub(r'[^\d]', '', candidate)
                    digits_only_full = re.sub(r'[^\d]', '', user_message)
                    # Accept only if the entire message consists of exactly 10 digits (no extra digits around)
                    if len(digits_only_full) == 10 and len(digits_only_match) == 10 and re.match(r'^[6-9]', digits_only_match):
                        phone = candidate
                        break
            
            # If no pattern matched, check if it's a valid Indian phone number (exactly 10 digits starting with 6-9)
            if not phone and re.match(r'^[\d\s\-\(\)\+]+$', user_message.strip()):
                digits_only = re.sub(r'[^\d]', '', user_message)
                if len(digits_only) == 10 and re.match(r'^[6-9]', digits_only):
                    phone = user_message.strip()
                else:
                    return "Please enter a valid phone number."
            
            if phone:
                state['phone'] = _normalize_phone(phone) or phone
                state['step'] = 3
                return generate_lead_capture_message(username, 3)
            else:
                return "Please enter a valid phone number."
        
        elif current_step == 3:  # Email
            email_pattern = r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
            # Allow skip if phone was provided earlier
            if user_message.strip().lower() == 'skip':
                if state.get('phone'):
                    state['email'] = ''
                    return complete_lead_capture(username, state, session_key, bot_id)
                else:
                    return "Please share at least one contact method (email or phone). What's your email address?"
            email_match = re.search(email_pattern, user_message)
            
            if email_match:
                email = email_match.group(1)
                state['email'] = email
                # Complete the lead capture
                return complete_lead_capture(username, state, session_key, bot_id)
            else:
                # If phone exists, allow finishing with just phone when user expresses no email
                if state.get('phone'):
                    state['email'] = ''
                    return complete_lead_capture(username, state, session_key, bot_id)
                return "I didn't catch your email address. Could you please provide your email? You can also type 'skip'."
        
        return None
        
    except Exception as e:
        logger.error(f"Error handling lead capture step: {e}")
        return "Thank you for your information! I've saved it and will be in touch."

def complete_lead_capture(username, state, session_key, bot_id=''):
    """Complete the lead capture process and save to database"""
    try:
        # Final server-side phone validation to avoid storing/displaying invalid numbers
        try:
            raw_phone = state.get('phone', '') or ''
            normalized = _normalize_phone(raw_phone)
            validated_phone = normalized if normalized else ''
        except Exception:
            validated_phone = ''

        # Save to Firebase
        if rtdb_available:
            lead_id = str(uuid.uuid4())
            assigned_to = assign_lead_round_robin(username)
            lead_data = {
                'id': lead_id,
                'name': state['name'],
                'email': state['email'],
                'phone': validated_phone,
                'source': 'chatbot',
                'created_at': int(time.time() * 1000),
                'username': username,
                'status': 'new',
                'message': '',
                'session_id': (session_key.split('_', 1)[1] if '_' in session_key else session_key),
                'assigned_to': assigned_to
            }
            base = _base_path(username, bot_id)
            fb_db.reference(f'{base}/leads').child(lead_id).set(lead_data)
            # Retroactively tag existing conversations in this session with the contact name
            try:
                sess = lead_data['session_id']
                cref = fb_db.reference(f'{base}/conversations')
                craw = cref.get() or {}
                if isinstance(craw, dict):
                    for cid, cval in craw.items():
                        if isinstance(cval, dict) and cval.get('session_id') == sess:
                            cref.child(cid).update({'contact_name': lead_data['name']})
            except Exception as _e:
                logger.warning(f"Failed to backfill contact_name on conversations: {_e}")
        
        # Generate completion response based on personality
        company_config = get_company_config(username)
        tone = company_config.get('tone', 'Professional') if company_config else 'Professional'
        company_name = company_config.get('companyName', 'our company') if company_config else 'our company'
        
        # Compose contact snippet conditionally (omit phone if invalid)
        contact_bits = []
        if validated_phone:
            contact_bits.append(f"Phone: {validated_phone}")
        if state.get('email'):
            contact_bits.append(f"Email: {state['email']}")
        contact_joined = ', '.join(contact_bits) if contact_bits else 'no contact details provided'

        # Tone-aware follow up
        tone_followups = {
            'Friendly': 'Anything else I can help you with? 😊',
            'Humorous': 'Anything else I can help you with before I power down? 😄',
            'Expert': 'Is there anything else I can assist you with?',
            'Caring': 'Is there anything else I can help you with? I’m here for you.',
            'Enthusiastic': 'Anything else I can help you with? 🚀',
            'Formal': 'Is there anything else with which I may assist you?',
            'Casual': 'Need anything else?'
        }
        followup = tone_followups.get(tone, 'Is there anything else I can help you with?')

        responses = {
            'Professional': f"Thank you, {state['name']}! I've saved your contact information ({contact_joined}). I'll make sure to reach out to you with relevant updates about {company_name}. {followup}",
            'Friendly': f"Awesome, {state['name']}! 😊 I've got your details saved ({contact_joined}). I'll keep you in the loop about all the exciting things happening at {company_name}! {followup}",
            'Humorous': f"Perfect! Thanks {state['name']}! 😄 I've added you to our VIP list at {company_name} ({contact_joined}). {followup}",
            'Expert': f"Excellent, {state['name']}. I've recorded your contact information ({contact_joined}) and will ensure you receive the most relevant updates about {company_name}. {followup}",
            'Caring': f"Thank you so much, {state['name']}! 🤗 I've saved your information ({contact_joined}) and will make sure to take good care of you with updates about {company_name}. {followup}",
            'Enthusiastic': f"Fantastic, {state['name']}! 🚀 I'm so excited to have your contact info ({contact_joined})! You'll be the first to know about all the amazing things at {company_name}! {followup}",
            'Formal': f"Thank you, {state['name']}. I have recorded your contact information ({contact_joined}) and will ensure you receive appropriate communications regarding {company_name}. {followup}",
            'Casual': f"Cool, {state['name']}! I've got your info saved ({contact_joined}). I'll make sure to keep you updated about {company_name}. {followup}"
        }
        
        # Clear the lead capture state and mark as completed
        if session_key in lead_capture_state:
            del lead_capture_state[session_key]
        
        # Mark this session as having completed lead capture
        lead_capture_completed[session_key] = True
        logger.info(f"Lead capture completed for {session_key}")
        
        return responses.get(tone, responses['Professional'])
        
    except Exception as e:
        logger.error(f"Error completing lead capture: {e}")
        return "Thank you for sharing your information! I've saved it and will be in touch."

def handle_lead_information(user_message, username):
    """Handle when user provides lead information"""
    try:
        # Extract lead information from the message
        import re
        
        # Extract name
        name_patterns = [
            r'name[:\s]+([a-zA-Z\s]+)',
            r'my name is ([a-zA-Z\s]+)',
            r'i\'m ([a-zA-Z\s]+)',
            r'call me ([a-zA-Z\s]+)'
        ]
        name = None
        for pattern in name_patterns:
            match = re.search(pattern, user_message, re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                break
        
        # Extract email
        email_pattern = r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
        email_match = re.search(email_pattern, user_message)
        email = email_match.group(1) if email_match else None
        
        # Extract phone with strict Indian mobile number validation
        phone_patterns = [
            r'phone[:\s]+([+\d\s\-\(\)]+)',
            r'call[:\s]+([+\d\s\-\(\)]+)',
            r'(\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})'
        ]
        phone = None
        for pattern in phone_patterns:
            match = re.search(pattern, user_message)
            if match:
                candidate = match.group(1).strip()
                digits_only_match = re.sub(r'[^\d]', '', candidate)
                digits_only_full = re.sub(r'[^\d]', '', user_message)
                if len(digits_only_full) == 10 and len(digits_only_match) == 10 and re.match(r'^[6-9]', digits_only_match):
                    phone = candidate
                    break
        
        # If no pattern matched, check if it's a valid Indian phone number (exactly 10 digits starting with 6-9)
        if not phone and re.match(r'^[\d\s\-\(\)\+]+$', user_message.strip()):
            digits_only = re.sub(r'[^\d]', '', user_message)
            if len(digits_only) == 10 and re.match(r'^[6-9]', digits_only):
                phone = user_message.strip()
            # If it's not a valid Indian number, don't extract it as phone
        
        # If we found lead information, create the lead
        if name and (email or phone):
            lead_data = {
                'name': name,
                'email': email or '',
                'phone': phone or '',
                'source': 'chatbot',
                'created_at': int(time.time() * 1000),
                'username': username,
                'status': 'new',
                'message': (message or '').strip() if (message := None) else ''
            }
            
            # Save to Firebase
            if rtdb_available:
                lead_id = str(uuid.uuid4())
                lead_data['id'] = lead_id
                lead_data['assigned_to'] = assign_lead_round_robin(username)
                fb_db.reference(f'{username}/leads').child(lead_id).set(lead_data)
            
            # Generate response based on personality
            company_config = get_company_config(username)
            tone = company_config.get('tone', 'Professional') if company_config else 'Professional'
            company_name = company_config.get('companyName', 'our company') if company_config else 'our company'
            
            responses = {
                'Professional': f"Thank you, {name}! I've saved your contact information. I'll make sure to reach out to you with relevant updates about {company_name}.",
                'Friendly': f"Awesome, {name}! 😊 I've got your details saved. I'll keep you in the loop about all the exciting things happening at {company_name}!",
                'Humorous': f"Perfect! Thanks {name}! 😄 I've added you to our VIP list at {company_name}. You're now officially part of the family!",
                'Expert': f"Excellent, {name}. I've recorded your contact information and will ensure you receive the most relevant updates about {company_name}.",
                'Caring': f"Thank you so much, {name}! 🤗 I've saved your information and will make sure to take good care of you with updates about {company_name}.",
                'Enthusiastic': f"Fantastic, {name}! 🚀 I'm so excited to have your contact info! You'll be the first to know about all the amazing things at {company_name}!",
                'Formal': f"Thank you, {name}. I have recorded your contact information and will ensure you receive appropriate communications regarding {company_name}.",
                'Casual': f"Cool, {name}! I've got your info saved. I'll make sure to keep you updated about {company_name}."
            }
            
            return responses.get(tone, responses['Professional'])
        else:
            # Ask for missing information
            missing = []
            if not name:
                missing.append("name")
            if not email and not phone:
                missing.append("email or phone number")
            
            return f"I'd love to save your contact information! Could you please provide your {' and '.join(missing)}?"
            
    except Exception as e:
        logger.error(f"Error handling lead information: {e}")
        return "Thank you for sharing your information! I've saved it and will be in touch."

# Add missing /chat endpoint for compatibility (some widgets may call this)
@app.route('/chat', methods=['POST'])
def chat():
    """Legacy endpoint for compatibility - redirects to send_message without strict auth"""
    try:
        # For testing purposes, allow chat without strict authentication
        user_message = request.json.get('message', '') if request.is_json else ''
        username = (request.json.get('username') or '').strip() if request.is_json else ''
        session_id = request.json.get('session_id', 'default') if request.is_json else 'default'
        
        if not user_message:
            return jsonify({'response': 'Please provide a message'}), 400
        
        # Set a default current_user for compatibility
        request.current_user = username
        
        # Load knowledge base for this user to ensure proper responses
        load_knowledge_base_for(username)
        
        # Get the regular response using the same logic as send_message
        bot_response = get_chatgpt_response(user_message)
        
        # Store conversation in Firebase (optional)
        try:
            if rtdb_available:
                conversation_id = str(uuid.uuid4())
                conversation_data = {
                    'id': conversation_id,
                    'user_message': user_message,
                    'bot_response': bot_response,
                    'timestamp': int(time.time() * 1000),
                    'session_id': session_id,
                    'username': username
                }
                fb_db.reference(f'{username}/conversations').child(conversation_id).set(conversation_data)
        except Exception as e:
            logger.warning(f'Failed to save conversation to RTDB: {e}')
        
        return jsonify({'response': bot_response})
    except Exception as e:
        logger.error(f'Error in /chat endpoint: {str(e)}')
        return jsonify({'response': 'I apologize for the inconvenience, but I am currently experiencing some technical difficulties. Please try again in a moment.'}), 500

@app.route('/send_message', methods=['POST'])
@require_auth_lenient
def send_message():
    try:
        user_message = request.json['message']
        username = (request.json.get('username') or '').strip()
        bot_id = (request.json.get('bot_id') or '').strip()
        session_id = request.json.get('session_id', 'default')
        
        # Check for appointment-related keywords
        appointment_keywords = ['schedule', 'appointment', 'book', 'meeting', 'consultation']
        cancel_keywords = ['cancel', 'reschedule', 'change appointment']
        
        msg_lower = user_message.lower()
        
        # Handle appointment cancellation by ID FIRST to avoid being caught by other keywords
        if msg_lower.startswith('apt-') or 'appointment id:' in msg_lower:
            bot_response = handle_appointment_cancellation(user_message, username)
        # Handle appointment scheduling
        elif any(keyword in msg_lower for keyword in appointment_keywords):
            appointment_state['waiting_for_title'] = True
            bot_response = "I'd be happy to help you schedule an appointment! First, please tell me the title of your appointment."
        # Handle appointment cancellation intent
        elif any(keyword in msg_lower for keyword in cancel_keywords):
            bot_response = "I can help you cancel an appointment. Please provide your appointment ID. If you don't have it, I can help you find your appointments."
        # Check if user is providing appointment title (step 1)
        elif appointment_state.get('waiting_for_title', False):
            bot_response = handle_appointment_title(user_message)
        # Check if user is providing appointment details
        elif 'title:' in msg_lower and ('date' in msg_lower or 'time' in msg_lower):
            bot_response = handle_appointment_scheduling(user_message)
        # Check if user wants to view appointments/calendar
        elif any(keyword in msg_lower for keyword in ['view appointments', 'show appointments', 'my appointments', 'calendar', 'appointment list']):
            bot_response = "I'll show you your appointments. You can view all appointments or search by appointment ID."
        # (moved above) Check if user is providing appointment ID for cancellation
        else:
            # Check if we're in lead capture mode first
            session_key = f"{username}_{session_id}"
            if session_key in lead_capture_state:
                # We're in lead capture mode, handle the response
                lead_response = handle_lead_capture_step(user_message, username, session_key, bot_id)
                if lead_response:
                    bot_response = lead_response
                else:
                    # If lead capture didn't handle it, get regular response
                    bot_response = get_chatgpt_response(user_message)
            else:
                # Get the regular response
                bot_response = get_chatgpt_response(user_message)
                
                # Increment message count first
                increment_message_count(username, session_id)
                
                # Check if we should ask for lead information after 2 messages
                message_count = get_message_count_simple(username, session_id)
                has_contact_info = has_user_provided_contact_info(username, session_id)
                in_lead_capture = session_key in lead_capture_state
                
                logger.info(f"Lead capture check for {username}: message_count={message_count}, has_contact_info={has_contact_info}, in_lead_capture={in_lead_capture}")
                
                if message_count >= 2 and not has_contact_info and not in_lead_capture:
                    # Start lead capture process
                    logger.info(f"Starting lead capture for {username}")
                    lead_capture_state[session_key] = {'step': 1, 'name': '', 'email': '', 'phone': ''}
                    lead_capture_msg = generate_lead_capture_message(username, 1)
                    if lead_capture_msg:
                        # Make it conversational by integrating with the response
                        bot_response = bot_response + " " + lead_capture_msg
        
        # Store conversation in Firebase (store under contact name when available, else session id)
        try:
            if rtdb_available:
                conversation_id = str(uuid.uuid4())
                # Try to resolve contact name by session lead
                contact_name = ''
                try:
                    base = _base_path(username, bot_id)
                    leads_index = fb_db.reference(f'{base}/leads').order_by_child('session_id').equal_to(session_id).get() or {}
                    if isinstance(leads_index, dict):
                        for _lid, _ldata in leads_index.items():
                            if isinstance(_ldata, dict) and _ldata.get('name'):
                                contact_name = str(_ldata.get('name'))
                                break
                except Exception as _:
                    pass
                display_key = contact_name or ''
                conversation_data = {
                    'id': conversation_id,
                    'user_message': user_message,
                    'bot_response': bot_response,
                    'timestamp': int(time.time() * 1000),
                    'session_id': request.json.get('session_id', 'default'),
                    'username': username or 'anonymous',
                    'bot_id': bot_id,
                    'contact_name': contact_name,
                    'display_key': display_key
                }
                base = _base_path(username, bot_id)
                fb_db.reference(f'{base}/conversations').child(conversation_id).set(conversation_data)
        except Exception as e:
            logger.warning(f"Failed to save conversation to RTDB: {e}")
        
        return jsonify({'response': bot_response})
    except Exception as e:
        logger.error(f"Error processing message: {str(e)}")
        return jsonify({'response': f"Error: {str(e)}"}), 500

@app.route('/schedule_appointment', methods=['POST'])
@require_auth_lenient
def schedule_appointment():
    try:
        if not rtdb_available:
            return jsonify({'error': 'Appointment storage is not configured.'}), 503

        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        title = data.get('title')
        appointment_time_str = data.get('time')
        username = (data.get('username') or '').strip()
        bot_id = (data.get('bot_id') or '').strip()
        session_id = (data.get('session_id') or '').strip()
        contact_name = (data.get('contact_name') or '').strip()

        if not title or not appointment_time_str:
            return jsonify({'error': 'Missing required fields: title and time are required'}), 400

        # Convert time string to datetime object with better error handling
        try:
            # Handle different time formats
            if appointment_time_str.endswith('Z'):
                appointment_time = datetime.fromisoformat(appointment_time_str.replace('Z', '+00:00'))
            elif '+' in appointment_time_str or appointment_time_str.endswith('00:00'):
                appointment_time = datetime.fromisoformat(appointment_time_str)
            else:
                # Assume local time if no timezone info
                appointment_time = datetime.fromisoformat(appointment_time_str)
                # Convert to UTC for storage
                appointment_time = appointment_time.replace(tzinfo=timezone.utc)
        except ValueError as e:
            logger.error(f"Invalid time format: {appointment_time_str}, error: {e}")
            return jsonify({'error': f'Invalid time format: {appointment_time_str}. Please use ISO format.'}), 400

        # Canonical slot key (UTC start-of-slot) for coarse locking
        appointment_utc = appointment_time.astimezone(timezone.utc)
        slot_key = appointment_utc.strftime('%Y%m%d-%H%M')  # e.g., 20250829-1630

        # Ensure the slot is not already booked (status not cancelled) via two checks:
        # 1) Query by exact time
        try:
            time_key = appointment_time.isoformat()
            base_path = _base_path(username, bot_id)
            existing = fb_db.reference(f'{base_path}/appointments').order_by_child('time').equal_to(time_key).get() or {}
            if isinstance(existing, dict):
                for appt in existing.values():
                    if appt and isinstance(appt, dict) and appt.get('status', 'pending') != 'cancelled':
                        return jsonify({'error': 'This time slot is already booked. Please choose a different slot.'}), 409
        except Exception as e:
            logger.warning(f"Slot availability check failed: {e}")

        # 2) Lock path to avoid race conditions
        try:
            base_path = _base_path(username, bot_id)
            lock_ref = fb_db.reference(f'{base_path}/slot_locks').child(slot_key)
            lock = lock_ref.get()
            if lock and isinstance(lock, dict) and lock.get('status') == 'booked':
                return jsonify({'error': 'This time slot is already booked. Please choose a different slot.'}), 409
        except Exception as e:
            logger.warning(f"Slot lock read failed: {e}")

        # Generate a unique ID
        appointment_id = f"APT-{int(time.time())}-{uuid.uuid4().hex[:8]}"
        
        # Create appointment object
        appointment = {
            'id': appointment_id,
            'title': title,
            'time': appointment_time.isoformat(),
            'status': 'pending',
            'created_at': int(time.time() * 1000),
            'username': username or 'anonymous',
            'bot_id': bot_id,
            'session_id': session_id,
            'contact_name': contact_name,
            'assigned_to': assign_appointment_round_robin(username or '')
        }
        
        # Save to Firebase Realtime Database
        try:
            base_path = _base_path(username, bot_id)
            fb_db.reference(f'{base_path}/appointments').child(appointment_id).set(appointment)
            logger.info(f"Appointment saved successfully: {appointment_id}")
        except Exception as e:
            logger.error(f"Failed to save appointment to Firebase: {e}")
            return jsonify({'error': 'Failed to save appointment. Please try again.'}), 500
            
        # Create/Update lock (scoped under user)
        try:
            base_path = _base_path(username, bot_id)
            fb_db.reference(f'{base_path}/slot_locks').child(slot_key).set({'status': 'booked', 'appointment_id': appointment_id})
            logger.info(f"Slot lock created successfully: {slot_key}")
        except Exception as e:
            logger.warning(f"Failed to set slot lock: {e}")
        
        return jsonify({
            'success': True,
            'message': 'Appointment scheduled successfully',
            'appointment': appointment,
            'appointment_id': appointment_id
        })
    except Exception as e:
        logger.error(f"Error scheduling appointment: {str(e)}")
        return jsonify({'error': f'Failed to schedule appointment: {str(e)}'}), 500

@app.route('/get_appointments', methods=['GET'])
@require_auth
def get_appointments():
    try:
        if not rtdb_available:
            return jsonify({'appointments': []})

        appointment_id = request.args.get('appointment_id')
        
        if appointment_id:
            # Get specific appointment by ID
            appointment = fb_db.reference('appointments').child(appointment_id).get()
            if appointment:
                return jsonify({'appointments': [appointment]})
            else:
                return jsonify({'appointments': []})
        else:
            # Get all appointments (scoped by bot if provided)
            user = (request.args.get('username') or '').strip() or 'anonymous'
            bot = (request.args.get('botId') or '').strip()
            base = _base_path(user, bot)
            snapshot = fb_db.reference(f'{base}/appointments').get() or {}
            appointments_list = list(snapshot.values()) if isinstance(snapshot, dict) else []
            
            # Sort by creation date (newest first)
            appointments_list.sort(key=lambda x: x.get('created_at', 0), reverse=True)
            
        return jsonify({'appointments': appointments_list})
    except Exception as e:
        logger.error(f"Error getting appointments: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/get_slot_locks', methods=['GET'])
@require_auth_lenient
def get_slot_locks():
    """Return slot lock statuses for a given date (YYYY-MM-DD).
    Response: { locks: { "YYYYMMDD-HHMM": "booked" | "cancelled" } }
    """
    try:
        if not rtdb_available:
            return jsonify({'locks': {}})

        date_str = request.args.get('date', '').strip()
        if not date_str:
            return jsonify({'locks': {}}), 400

        # Fetch all locks (small dataset expected); filter by date prefix on server side
        user = (request.args.get('username') or '').strip() or 'anonymous'
        bot = (request.args.get('botId') or '').strip()
        base = _base_path(user, bot)
        all_locks = fb_db.reference(f'{base}/slot_locks').get() or {}
        locks = {}
        try:
            # Build prefix for keys like YYYYMMDD-
            y, m, d = map(int, date_str.split('-'))
            prefix = f"{y}{str(m).zfill(2)}{str(d).zfill(2)}-"
            if isinstance(all_locks, dict):
                for k, v in all_locks.items():
                    if isinstance(k, str) and k.startswith(prefix):
                        if isinstance(v, dict):
                            status = v.get('status', 'booked')
                        else:
                            status = str(v)
                        locks[k] = status
        except Exception as e:
            logger.warning(f"Failed to filter slot locks: {e}")
        return jsonify({'locks': locks})
    except Exception as e:
        logger.error(f"Error getting slot locks: {str(e)}")
        return jsonify({'locks': {}}), 500

@app.route('/cancel_appointment', methods=['POST'])
@require_auth_lenient
def cancel_appointment():
    try:
        if not rtdb_available:
            return jsonify({'error': 'Appointment storage is not configured.'}), 503

        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        appointment_id = data.get('appointment_id')
        username = (data.get('username') or '').strip()
        bot_id = (data.get('bot_id') or '').strip()

        if not appointment_id:
            return jsonify({'error': 'Appointment ID is required'}), 400

        # Get the appointment first to check if it exists
        base_path = _base_path(username, bot_id)
        appt_ref = fb_db.reference(f'{base_path}/appointments').child(appointment_id)
        appt = appt_ref.get()
        
        if not appt:
            return jsonify({'error': 'Appointment not found'}), 404
            
        if appt.get('status') == 'cancelled':
            return jsonify({'error': 'Appointment is already cancelled'}), 400

        # Update appointment status in Firebase and free lock
        try:
            appt_ref.update({
                'status': 'cancelled', 
                'cancelled_by': username or 'anonymous',
                'cancelled_at': int(time.time() * 1000)
            })
            logger.info(f"Appointment cancelled successfully: {appointment_id}")
        except Exception as e:
            logger.error(f"Failed to update appointment status: {e}")
            return jsonify({'error': 'Failed to cancel appointment. Please try again.'}), 500
            
        # Free the slot lock
        try:
            appt_time = appt.get('time')
            if appt_time:
                appt_dt = datetime.fromisoformat(appt_time.replace('Z', '+00:00')).astimezone(timezone.utc)
                slot_key = appt_dt.strftime('%Y%m%d-%H%M')
                fb_db.reference(f'{base_path}/slot_locks').child(slot_key).set({'status': 'cancelled', 'appointment_id': appointment_id})
                logger.info(f"Slot lock freed successfully: {slot_key}")
        except Exception as e:
            logger.warning(f"Failed to free slot lock: {e}")

        return jsonify({
            'success': True,
            'message': 'Appointment cancelled successfully',
            'appointment_id': appointment_id
        })
    except Exception as e:
        logger.error(f"Error cancelling appointment: {str(e)}")
        return jsonify({'error': f'Failed to cancel appointment: {str(e)}'}), 500

@app.route('/create_lead', methods=['POST'])
@require_auth_lenient
def create_lead():
    try:
        if not rtdb_available:
            return jsonify({'success': False, 'message': 'Lead storage is not configured.'}), 503

        data = request.json or {}
        name = (data.get('name') or '').strip()
        email = (data.get('email') or '').strip()
        phone = (data.get('phone') or '').strip()
        message = (data.get('message') or '').strip()
        username = (data.get('username') or '').strip() or (request.headers.get('X-User') or '').strip()
        session_id = (data.get('session_id') or '').strip()
        bot_id = (data.get('bot_id') or '').strip()

        if not name or not (email or phone):
            return jsonify({'success': False, 'message': 'Please provide your name and at least one contact (email or phone).'}), 400

        lead_id = str(uuid.uuid4())
        lead_data = {
            'id': lead_id,
            'name': name,
            'email': email,
            'phone': phone,
            'message': message,
            'source': 'chatbot',
            'created_at': int(time.time() * 1000),
            'username': username or 'anonymous',
            'status': 'new',
            'session_id': session_id,
            'assigned_to': assign_lead_round_robin(username)
        }

        base = _base_path(username, bot_id)
        fb_db.reference(f'{base}/leads').child(lead_id).set(lead_data)
        return jsonify({'success': True, 'message': 'Thanks! Your contact has been shared successfully.', 'lead_id': lead_id})
    except Exception as e:
        logger.error(f"Error creating lead: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/save_company_config', methods=['POST'])
def save_company_config():
    try:
        if not rtdb_available:
            return jsonify({'success': False, 'message': 'Database not configured.'}), 503

        data = request.json or {}
        username = (data.get('username') or '').strip()
        
        if not username:
            return jsonify({'success': False, 'message': 'Missing username'}), 400

        # Extract company configuration data
        company_config = {
            'companyName': data.get('companyName', ''),
            'companyUrl': data.get('companyUrl', ''),
            'companyDescription': data.get('companyDescription', ''),
            'primaryColor': data.get('primaryColor', '#4f46e5'),
            'tone': data.get('tone', 'Professional'),
            'industry': data.get('industry', ''),
            'welcomeMessage': data.get('welcomeMessage', ''),
            'avatarUrl': data.get('avatarUrl', ''),
            'files': data.get('files', []),
            'createdAt': datetime.utcnow().isoformat() + 'Z',
            'updatedAt': datetime.utcnow().isoformat() + 'Z'
        }

        # Save to Firebase
        fb_db.reference(f'{username}/company_config').set(company_config)
        
        return jsonify({'success': True, 'message': 'Company configuration saved successfully!'})
    except Exception as e:
        logger.error(f"Error saving company config: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/upload_knowledge_base', methods=['POST'])
@require_auth
def upload_knowledge_base():
    try:
        if not rtdb_available:
            return jsonify({'success': False, 'message': 'Database not configured.'}), 503

        # Check if file is uploaded
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'No file uploaded.'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'No file selected.'}), 400

        if file:
            try:
                username = (request.args.get('username') or request.form.get('username') or request.headers.get('X-User') or '').strip()
                if not username:
                    return jsonify({'success': False, 'message': 'Missing username'}), 400
                
                # Get file type
                file_type = request.form.get('fileType', '')
                file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
                
                # Helper to decode text safely
                def _read_text_safe(file_obj) -> str:
                    try:
                        return file_obj.read().decode('utf-8')
                    except Exception:
                        try:
                            file_obj.seek(0)
                        except Exception:
                            pass
                        return file_obj.read().decode('latin-1', errors='replace')

                # CSV/TSV parser
                def _parse_csv(text: str):
                    sample = text[:4096]
                    try:
                        dialect = csv.Sniffer().sniff(sample)
                        delimiter = dialect.delimiter
                    except Exception:
                        delimiter = '\t' if '\t' in sample and ',' not in sample else ','
                    reader = csv.DictReader(io.StringIO(text), delimiter=delimiter)
                    rows = [dict(r) for r in reader]
                    return {
                        "type": "table",
                        "columns": reader.fieldnames or [],
                        "rows": rows
                    }

                # XLSX parser (optional)
                def _parse_xlsx(stream_bytes: bytes):
                    try:
                        from openpyxl import load_workbook  # type: ignore
                    except Exception:
                        return {
                            "type": "spreadsheet",
                            "content": stream_bytes.hex(),
                            "note": "openpyxl not installed; stored raw"
                        }
                    wb = load_workbook(filename=io.BytesIO(stream_bytes), read_only=True, data_only=True)
                    sheet = wb.active
                    rows = []
                    headers = []
                    for idx, row in enumerate(sheet.iter_rows(values_only=True)):
                        values = ["" if v is None else v for v in list(row)]
                        if idx == 0:
                            headers = [str(h) for h in values]
                        else:
                            item = {}
                            for j, h in enumerate(headers or []):
                                item[str(h)] = str(values[j]) if j < len(values) else ""
                            if not headers:
                                item = {str(j): str(v) for j, v in enumerate(values)}
                            rows.append(item)
                    return {
                        "type": "table",
                        "columns": headers,
                        "rows": rows
                    }

                # PDF parser (optional)
                def _parse_pdf(stream_bytes: bytes):
                    try:
                        from PyPDF2 import PdfReader  # type: ignore
                    except Exception:
                        return {"type": "document", "content": stream_bytes.hex(), "note": "PyPDF2 not installed; stored raw"}
                    text = []
                    reader = PdfReader(io.BytesIO(stream_bytes))
                    for page in reader.pages:
                        try:
                            text.append(page.extract_text() or "")
                        except Exception:
                            pass
                    return {"type": "text", "content": "\n\n".join(text)}

                # DOCX parser (optional)
                def _parse_docx(stream_bytes: bytes):
                    try:
                        import docx  # type: ignore
                    except Exception:
                        return {"type": "document", "content": stream_bytes.hex(), "note": "python-docx not installed; stored raw"}
                    document = docx.Document(io.BytesIO(stream_bytes))
                    parts = [p.text for p in document.paragraphs if p.text]
                    # Include tables
                    for table in document.tables:
                        for row in table.rows:
                            parts.append(" | ".join([cell.text for cell in row.cells]))
                    return {"type": "text", "content": "\n".join(parts)}

                # Process different file types
                if file_extension == 'json' or file_type == 'application/json':
                    # Handle JSON files
                    content = _read_text_safe(file)
                    knowledge_data = json.loads(content)
                    
                    if isinstance(knowledge_data, list):
                        knowledge_data = {"data": knowledge_data}
                    elif not isinstance(knowledge_data, dict):
                        knowledge_data = {"content": str(knowledge_data)}
                        
                elif file_extension in ['csv', 'tsv'] or ('text' in file_type and ('csv' in file.filename.lower() or 'tsv' in file.filename.lower())):
                    # Handle CSV/TSV files
                    text = _read_text_safe(file)
                    table = _parse_csv(text)
                    knowledge_data = {"filename": file.filename, **table}

                elif file_extension in ['txt', 'md'] or ('text' in file_type):
                    # Handle generic text files
                    content = _read_text_safe(file)
                    knowledge_data = {"filename": file.filename, "content": content, "type": "text"}
                    
                elif file_extension in ['pdf', 'doc', 'docx'] or 'pdf' in file_type or 'word' in file_type:
                    # Handle PDF and Word files (store as binary for now)
                    blob = file.read()
                    if file_extension == 'pdf' or 'pdf' in file_type:
                        knowledge_data = {"filename": file.filename, **_parse_pdf(blob)}
                    elif file_extension in ['doc', 'docx'] or 'word' in file_type:
                        knowledge_data = {"filename": file.filename, **_parse_docx(blob)}
                    else:
                        knowledge_data = {"filename": file.filename, "content": blob.hex(), "type": "document", "original_type": file_type}
                    
                elif file_extension in ['xls', 'xlsx'] or 'excel' in file_type or 'spreadsheet' in file_type:
                    # Handle Excel files
                    blob = file.read()
                    knowledge_data = {"filename": file.filename, **_parse_xlsx(blob)}
                else:
                    return jsonify({'success': False, 'message': f'Unsupported file type: {file_extension}'}), 400
                
                # Store file metadata in user's knowledge base
                file_id = f"file_{int(time.time())}_{file.filename.replace(' ', '_')}"
                knowledge_data['file_id'] = file_id
                knowledge_data['uploaded_at'] = datetime.utcnow().isoformat() + 'Z'
                
                # Update per-user knowledge base
                if update_knowledge_base(knowledge_data, username):
                    return jsonify({'success': True, 'message': f'File {file.filename} uploaded successfully!'})
                else:
                    return jsonify({'success': False, 'message': 'Failed to update knowledge base.'}), 500
                    
            except json.JSONDecodeError as e:
                return jsonify({'success': False, 'message': f'Invalid JSON format: {str(e)}'}), 400
            except Exception as e:
                return jsonify({'success': False, 'message': f'Error processing file: {str(e)}'}), 500
        else:
            return jsonify({'success': False, 'message': 'Please upload a valid file.'}), 400

    except Exception as e:
        logger.error(f"Error uploading knowledge base: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/get_company_config', methods=['GET'])
def get_company_config_endpoint():
    try:
        username = (request.args.get('username') or '').strip()
        if not username:
            return jsonify({'success': False, 'message': 'Missing username'}), 400
        
        config = get_company_config(username)
        if not config:
            return jsonify({'success': False, 'message': 'Company configuration not found'}), 404
        
        return jsonify({'success': True, 'config': config})
    except Exception as e:
        logger.error(f"Error getting company config: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/get_knowledge_base', methods=['GET'])
@require_auth
def get_knowledge_base():
    try:
        username = (request.args.get('username') or '').strip()
        data = KB_CACHE.get(username) or load_knowledge_base_for(username)
        formatted_data = format_knowledge_base(data or {})
        return jsonify({
            'success': True, 
            'raw_data': data or {},
            'formatted_data': formatted_data
        })
    except Exception as e:
        logger.error(f"Error getting knowledge base: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/reload_knowledge_base', methods=['POST'])
@require_auth
def reload_knowledge_base():
    try:
        username = (request.args.get('username') or '').strip()
        load_knowledge_base_for(username)
        return jsonify({'success': True, 'message': 'Knowledge base reloaded successfully!', 'username': username})
    except Exception as e:
        logger.error(f"Error reloading knowledge base: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

# --- Simple auth endpoints ---
def _issue_token(username: str) -> str:
    payload = {
        'sub': username,
        'iat': int(time.time()),
        'exp': int(time.time()) + JWT_EXP_SECONDS
    }
    return jwt.encode(payload, JWT_SECRET, algorithm='HS256')

def _verify_token(token: str) -> str:
    try:
        data = jwt.decode(token, JWT_SECRET, algorithms=['HS256'])
        return str(data.get('sub') or '')
    except Exception:
        return ''

def _otp_hash(code: str) -> str:
    try:
        base = (code or '').strip()
        return hashlib.sha256((base + str(JWT_SECRET)).encode('utf-8')).hexdigest()
    except Exception:
        return ''

def _twilio_client():
    try:
        from twilio.rest import Client  # type: ignore
        sid = os.getenv('TWILIO_ACCOUNT_SID')
        token = os.getenv('TWILIO_AUTH_TOKEN')
        if not sid or not token:
            return None
        return Client(sid, token)
    except Exception as _e:
        logger.warning(f"Twilio client unavailable: {_e}")
        return None

@app.route('/auth/send_otp', methods=['POST'])
def auth_send_otp():
    try:
        data = request.json or {}
        username = (data.get('username') or '').strip()
        phone = (data.get('phone') or '').strip()
        if not username or not phone:
            return jsonify({'success': False, 'message': 'Username and phone are required'}), 400

        # Normalize phone (basic)
        phone_norm = phone if phone.startswith('+') else ('+91' + ''.join([c for c in phone if c.isdigit()]))

        otp_code = str(random.randint(100000, 999999))
        expires_at = int(time.time()) + 5 * 60

        # Store OTP in DB (hashed)
        if rtdb_available:
            otp_ref = fb_db.reference('otp_codes').child(username)
            otp_ref.set({
                'phone': phone_norm,
                'code_hash': _otp_hash(otp_code),
                'createdAt': datetime.utcnow().isoformat() + 'Z',
                'expiresAt': expires_at
            })
        else:
            return jsonify({'success': False, 'message': 'Database not available'}), 503

        # Send SMS via Twilio
        sent = False
        try:
            client = _twilio_client()
            from_number = os.getenv('TWILIO_NUMBER')
            messaging_service_sid = os.getenv('TWILIO_MESSAGING_SERVICE_SID')
            if client and (from_number or messaging_service_sid):
                create_kwargs = {
                    'body': f"Your IM Solutions verification code is {otp_code}. It expires in 5 minutes.",
                    'to': phone_norm,
                }
                if from_number:
                    create_kwargs['from_'] = from_number
                if messaging_service_sid and 'from_' not in create_kwargs:
                    create_kwargs['messaging_service_sid'] = messaging_service_sid
                client.messages.create(**create_kwargs)
                sent = True
        except Exception as e:
            logger.warning(f"Failed to send OTP SMS: {e}")

        return jsonify({'success': True, 'sent': sent})
    except Exception as e:
        logger.error(f"send_otp error: {e}")
        return jsonify({'success': False, 'message': 'Failed to send OTP'}), 500

@app.route('/auth/signup', methods=['POST'])
def auth_signup():
    try:
        data = request.json or {}
        username = (data.get('username') or '').strip()
        email = (data.get('email') or '').strip()
        password = (data.get('password') or '').strip()
        confirm_password = (data.get('confirm_password') or '').strip()
        otp_code = (data.get('otp') or '').strip()
        phone = (data.get('phone') or '').strip()
        
        if not username or not email or not password:
            return jsonify({'success': False, 'message': 'Missing username, email or password'}), 400
        
        if password != confirm_password:
            return jsonify({'success': False, 'message': 'Passwords do not match'}), 400
        
        if len(password) < 6:
            return jsonify({'success': False, 'message': 'Password must be at least 6 characters long'}), 400
        
        if len(username) < 3:
            return jsonify({'success': False, 'message': 'Username must be at least 3 characters long'}), 400
        
        # Basic email validation
        import re
        email_pattern = r'^[^\s@]+@[^\s@]+\.[^\s@]+$'
        if not re.match(email_pattern, email):
            return jsonify({'success': False, 'message': 'Please enter a valid email address'}), 400

        # Validate OTP if provided (required when phone supplied)
        if phone:
            if not otp_code:
                return jsonify({'success': False, 'message': 'OTP is required'}), 400
            if not rtdb_available:
                return jsonify({'success': False, 'message': 'Database not available'}), 503
            otp_ref = fb_db.reference('otp_codes').child(username)
            rec = otp_ref.get() or {}
            if not rec:
                return jsonify({'success': False, 'message': 'No OTP requested'}), 400
            if int(rec.get('expiresAt') or 0) < int(time.time()):
                return jsonify({'success': False, 'message': 'OTP expired'}), 400
            if _otp_hash(otp_code) != str(rec.get('code_hash') or ''):
                return jsonify({'success': False, 'message': 'Invalid OTP'}), 400

        # Check if user already exists
        try:
            if rtdb_available:
                user_ref = fb_db.reference('users').child(username)
                existing_user = user_ref.get()
                if existing_user:
                    return jsonify({'success': False, 'message': 'Username already exists'}), 409
                
                # Hash password and create user
                salt, password_hash = _hash_password(password)
                user_data = {
                    'id': username,
                    'username': username,
                    'email': email,
                    'phone': phone,
                    'password_hash': password_hash,
                    'salt': salt,
                    'role': 'user',
                    'createdAt': datetime.utcnow().isoformat() + 'Z'
                }
                user_ref.set(user_data)
                # clear OTP after sign up
                try:
                    fb_db.reference('otp_codes').child(username).delete()
                except Exception:
                    pass
            else:
                return jsonify({'success': False, 'message': 'Database not available'}), 503
                
        except Exception as e:
            logger.error(f"Signup error: {e}")
            return jsonify({'success': False, 'message': 'Registration failed'}), 500

        token = _issue_token(username)
        return jsonify({'success': True, 'token': token, 'username': username, 'message': 'Account created successfully'})
    except Exception as e:
        logger.error(f'Signup error: {e}')
        return jsonify({'success': False, 'message': 'Registration failed'}), 500

@app.route('/auth/login', methods=['POST'])
def auth_login():
    try:
        data = request.json or {}
        username = (data.get('username') or '').strip()
        password = (data.get('password') or '').strip()
        if not username or not password:
            return jsonify({'success': False, 'message': 'Missing username or password'}), 400

        # Check if user exists and verify password
        user = None
        try:
            if rtdb_available:
                user_ref = fb_db.reference('users').child(username)
                user = user_ref.get()
                if not user:
                    return jsonify({'success': False, 'message': 'User not found. Please sign up first.'}), 404
                
                # Verify password using stored salt and hash
                stored_salt = user.get('salt', '')
                stored_hash = user.get('password_hash', '')
                
                if not stored_salt or not stored_hash:
                    return jsonify({'success': False, 'message': 'Invalid user data. Please contact support.'}), 500
                
                if not _verify_password(password, stored_salt, stored_hash):
                    return jsonify({'success': False, 'message': 'Invalid credentials'}), 401
            else:
                return jsonify({'success': False, 'message': 'Database not available'}), 503
                
        except Exception as e:
            logger.error(f"Auth DB access failed: {e}")
            return jsonify({'success': False, 'message': 'Authentication failed'}), 500

        token = _issue_token(username)
        user_role = 'user'
        try:
            user_role = str((user or {}).get('role') or 'user')
        except Exception:
            user_role = 'user'
        return jsonify({'success': True, 'token': token, 'username': username, 'role': user_role})
    except Exception as e:
        logger.error(f'Auth error: {e}')
        return jsonify({'success': False, 'message': 'Authentication failed'}), 500

@app.route('/auth/verify', methods=['POST'])
def auth_verify():
    token = (request.json or {}).get('token') or ''
    user = _verify_token(token)
    if not user:
        return jsonify({'success': False}), 401
    return jsonify({'success': True, 'username': user})

@app.route('/logout')
def logout():
    """Logout route that redirects to login page"""
    return redirect('/')

# --- OAuth: Google ---
@app.route('/auth/google')
def auth_google_start():
    client_id = os.getenv('GOOGLE_CLIENT_ID')
    redirect_uri = (request.url_root.rstrip('/') + '/auth/google/callback')
    if not client_id:
        return jsonify({'success': False, 'message': 'Google OAuth not configured'}), 500
    params = {
        'client_id': client_id,
        'redirect_uri': redirect_uri,
        'response_type': 'code',
        'scope': 'openid email profile',
        'access_type': 'offline',
        'prompt': 'consent'
    }
    return redirect('https://accounts.google.com/o/oauth2/v2/auth?' + urlencode(params))

@app.route('/auth/google/callback')
def auth_google_callback():
    try:
        code = request.args.get('code')
        if not code:
            return redirect('/login')
        client_id = os.getenv('GOOGLE_CLIENT_ID')
        client_secret = os.getenv('GOOGLE_CLIENT_SECRET')
        redirect_uri = (request.url_root.rstrip('/') + '/auth/google/callback')
        token_res = requests.post('https://oauth2.googleapis.com/token', data={
            'code': code,
            'client_id': client_id,
            'client_secret': client_secret,
            'redirect_uri': redirect_uri,
            'grant_type': 'authorization_code'
        }, timeout=10)
        token_json = token_res.json()
        access_token = token_json.get('access_token')
        if not access_token:
            return redirect('/login')
        userinfo_res = requests.get('https://www.googleapis.com/oauth2/v3/userinfo', headers={
            'Authorization': f'Bearer {access_token}'
        }, timeout=10)
        info = userinfo_res.json() if userinfo_res.ok else {}
        email = (info.get('email') or '').strip()
        name = (info.get('name') or '').strip()
        if not email:
            return redirect('/login')
        username = email.replace('@', '_').replace('.', '_')
        # Upsert user in DB
        if rtdb_available:
            user_ref = fb_db.reference('users').child(username)
            existing = user_ref.get()
            if not existing:
                user_ref.set({
                    'id': username,
                    'username': username,
                    'email': email,
                    'name': name,
                    'provider': 'google',
                    'role': 'user',
                    'createdAt': datetime.utcnow().isoformat() + 'Z'
                })
        token = _issue_token(username)
        return redirect('/dashboard/' + username)
    except Exception as e:
        logger.error(f'Google OAuth error: {e}')
        return redirect('/login')

# --- OAuth: Facebook ---
@app.route('/auth/facebook')
def auth_facebook_start():
    client_id = os.getenv('FACEBOOK_CLIENT_ID')
    redirect_uri = (request.url_root.rstrip('/') + '/auth/facebook/callback')
    if not client_id:
        return jsonify({'success': False, 'message': 'Facebook OAuth not configured'}), 500
    params = {
        'client_id': client_id,
        'redirect_uri': redirect_uri,
        'response_type': 'code',
        'scope': 'email,public_profile'
    }
    return redirect('https://www.facebook.com/v17.0/dialog/oauth?' + urlencode(params))

@app.route('/auth/facebook/callback')
def auth_facebook_callback():
    try:
        code = request.args.get('code')
        if not code:
            return redirect('/login')
        client_id = os.getenv('FACEBOOK_CLIENT_ID')
        client_secret = os.getenv('FACEBOOK_CLIENT_SECRET')
        redirect_uri = (request.url_root.rstrip('/') + '/auth/facebook/callback')
        token_res = requests.get('https://graph.facebook.com/v17.0/oauth/access_token', params={
            'client_id': client_id,
            'client_secret': client_secret,
            'redirect_uri': redirect_uri,
            'code': code
        }, timeout=10)
        token_json = token_res.json()
        access_token = token_json.get('access_token')
        if not access_token:
            return redirect('/login')
        userinfo_res = requests.get('https://graph.facebook.com/me', params={
            'fields': 'id,name,email,picture',
            'access_token': access_token
        }, timeout=10)
        info = userinfo_res.json() if userinfo_res.ok else {}
        email = (info.get('email') or f"fb_{info.get('id','')}@example.com").strip()
        name = (info.get('name') or '').strip()
        username = email.replace('@', '_').replace('.', '_')
        if rtdb_available:
            user_ref = fb_db.reference('users').child(username)
            existing = user_ref.get()
            if not existing:
                user_ref.set({
                    'id': username,
                    'username': username,
                    'email': email,
                    'name': name,
                    'provider': 'facebook',
                    'role': 'user',
                    'createdAt': datetime.utcnow().isoformat() + 'Z'
                })
        token = _issue_token(username)
        return redirect('/dashboard/' + username)
    except Exception as e:
        logger.error(f'Facebook OAuth error: {e}')
        return redirect('/login')

@app.route('/debug/<username>')
def debug_data(username):
    """Debug route to check data structure"""
    token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    if not token:
        return jsonify({'error': 'No token provided'})
    
    verified_username = _verify_token(token)
    if not verified_username or verified_username != username:
        return jsonify({'error': 'Invalid token'})
    
    debug_info = {
        'username': username,
        'rtdb_available': rtdb_available,
        'data': {}
    }
    
    if rtdb_available:
        try:
            # Check bots
            bots_ref = fb_db.reference(f'{username}/bots')
            bots = bots_ref.get() or {}
            debug_info['data']['bots'] = bots
            
            # Check old structure
            leads_old = fb_db.reference(f'{username}/leads').get() or {}
            appointments_old = fb_db.reference(f'{username}/appointments').get() or {}
            conversations_old = fb_db.reference(f'{username}/conversations').get() or {}
            
            debug_info['data']['old_structure'] = {
                'leads_count': len(leads_old) if isinstance(leads_old, dict) else 0,
                'appointments_count': len(appointments_old) if isinstance(appointments_old, dict) else 0,
                'conversations_count': len(conversations_old) if isinstance(conversations_old, dict) else 0
            }
            
            # Check new structure if bot exists
            if bots and len(bots) > 0:
                bot_id = list(bots.keys())[0]
                base = _base_path(username, bot_id)
                
                leads_new = fb_db.reference(f'{base}/leads').get() or {}
                appointments_new = fb_db.reference(f'{base}/appointments').get() or {}
                conversations_new = fb_db.reference(f'{base}/conversations').get() or {}
                
                debug_info['data']['new_structure'] = {
                    'bot_id': bot_id,
                    'base_path': base,
                    'leads_count': len(leads_new) if isinstance(leads_new, dict) else 0,
                    'appointments_count': len(appointments_new) if isinstance(appointments_new, dict) else 0,
                    'conversations_count': len(conversations_new) if isinstance(conversations_new, dict) else 0
                }
        except Exception as e:
            debug_info['error'] = str(e)
    
    return jsonify(debug_info)

@app.route('/setup/<username>')
def setup_page(username):
    """Setup page for editing the user's single chatbot"""
    # Require token and ensure it matches the route username
    token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    verified_username = _verify_token(token) if token else ''
    if not verified_username or verified_username != username:
        return render_template('login.html')
    
    try:
        # Get user's single bot
        bot = None
        if rtdb_available:
            bots_ref = fb_db.reference(f'{username}/bots')
            bots = bots_ref.get() or {}
            if bots and len(bots) > 0:
                bot = list(bots.values())[0]
        
        # Get company config
        company_config = {}
        if rtdb_available:
            config_ref = fb_db.reference(f'{username}/company_config')
            company_config = config_ref.get() or {}
        
        # Get knowledge base
        knowledge_base = []
        if rtdb_available:
            kb_ref = fb_db.reference(f'{username}/knowledge_base')
            knowledge_base = kb_ref.get() or []
        
        return render_template('setup_knowledge.html', 
                             username=username, 
                             bot=bot, 
                             company_config=company_config,
                             knowledge_base=knowledge_base,
                             token=token)
    except Exception as e:
        logger.error(f"Setup page error: {e}")
        return render_template('setup_knowledge.html', 
                             username=username, 
                             bot=None, 
                             company_config={},
                             knowledge_base=[],
                             token=token)

@app.route('/setup')
def setup_page_token():
    token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    username = _verify_token(token) if token else ''
    if not username:
        return render_template('login.html')
    return setup_page(username)

# Sample data loader removed in production
def load_sample_knowledge_base():
    return

# Create default user account for testing
def create_default_user():
    """Create a default user account for testing (disabled: no hardcoded sample user)"""
    try:
        if rtdb_available:
            return jsonify({'success': False, 'message': 'Disabled'}), 403
            
            # Disabled
            user_ref = fb_db.reference('users').child(username)
            existing_user = user_ref.get()
            
            if not existing_user:
                # Hash password and create user
                salt, password_hash = _hash_password(password)
                user_data = {
                    'id': username,
                    'username': username,
                    'password_hash': password_hash,
                    'salt': salt,
                    'role': 'user',
                    'createdAt': datetime.utcnow().isoformat() + 'Z'
                }
                user_ref.set(user_data)
                logger.info(f"Default user created: {username}")
            else:
                logger.info(f"Default user already exists: {username}")
                
    except Exception as e:
        logger.error(f"Error creating default user: {e}")

@app.route('/users/<username>/subadmins', methods=['GET'])
@require_auth
def list_subadmins(username):
    try:
        # Verify the token subject matches the route username
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        subject = _verify_token(token)
        if not subject or subject != username:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 401

        if not rtdb_available:
            return jsonify({'success': False, 'message': 'Database not available'}), 503

        users_ref = fb_db.reference('users')
        # Query users by parent to avoid scanning entire set
        try:
            q = users_ref.order_by_child('parent').equal_to(username).get() or {}
        except Exception:
            # Fallback to full get if indexes are not set
            q = users_ref.get() or {}
        subadmins = []
        if isinstance(q, dict):
            for user_id, user_obj in q.items():
                try:
                    if (
                        isinstance(user_obj, dict)
                        and user_obj.get('role') == 'subadmin'
                        and user_obj.get('parent') == username
                    ):
                        subadmins.append({
                            'username': user_obj.get('username') or user_id,
                            'createdAt': user_obj.get('createdAt') or '',
                            'role': user_obj.get('role') or 'subadmin',
                            'access': user_obj.get('access') or {}
                        })
                except Exception:
                    continue
        return jsonify({'success': True, 'items': subadmins})
    except Exception as e:
        logger.error(f"List subadmins error: {e}")
        return jsonify({'success': False, 'message': 'Failed to list subadmins'}), 500


@app.route('/users/<username>/subadmins', methods=['POST'])
@require_auth
def create_subadmin(username):
    try:
        # Verify the token subject matches the route username (only main admin can create)
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        subject = _verify_token(token)
        if not subject or subject != username:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 401

        if not rtdb_available:
            return jsonify({'success': False, 'message': 'Database not available'}), 503

        data = request.json or {}
        sub_username = (data.get('username') or '').strip()
        password = (data.get('password') or '').strip()
        confirm_password = (data.get('confirm_password') or '').strip()
        # Access permissions
        access = data.get('access') or {}
        access_obj = {
            'leads': bool(access.get('leads', False)),
            'appointments': bool(access.get('appointments', False)),
            'conversations': bool(access.get('conversations', False))
        }

        if not sub_username or not password:
            return jsonify({'success': False, 'message': 'Missing username or password'}), 400
        if len(sub_username) < 3:
            return jsonify({'success': False, 'message': 'Username must be at least 3 characters long'}), 400
        if len(password) < 6:
            return jsonify({'success': False, 'message': 'Password must be at least 6 characters long'}), 400
        if password != confirm_password:
            return jsonify({'success': False, 'message': 'Passwords do not match'}), 400
        if sub_username == username:
            return jsonify({'success': False, 'message': 'Sub-admin username cannot be the same as main username'}), 400

        # Ensure the sub-admin user does not already exist
        user_ref = fb_db.reference('users').child(sub_username)
        existing = user_ref.get()
        if existing:
            return jsonify({'success': False, 'message': 'Username already exists'}), 409

        # Create sub-admin user in the global users directory with parent link
        salt, password_hash = _hash_password(password)
        now_iso = datetime.utcnow().isoformat() + 'Z'
        user_data = {
            'id': sub_username,
            'username': sub_username,
            'password_hash': password_hash,
            'salt': salt,
            'role': 'subadmin',
            'parent': username,
            'access': access_obj,
            'createdAt': now_iso
        }
        user_ref.set(user_data)
        return jsonify({'success': True, 'user': {'username': sub_username, 'role': 'subadmin', 'createdAt': now_iso, 'access': access_obj}})
    except Exception as e:
        logger.error(f"Create subadmin error: {e}")
        return jsonify({'success': False, 'message': 'Failed to create subadmin'}), 500

def _get_eligible_subadmins(username: str, access_key: str = 'leads') -> list:
    try:
        if not rtdb_available:
            return []
        all_users = fb_db.reference('users').get() or {}
        eligible = []
        if isinstance(all_users, dict):
            for uid, u in all_users.items():
                if not isinstance(u, dict):
                    continue
                if u.get('role') == 'subadmin' and u.get('parent') == username:
                    access = u.get('access') or {}
                    if bool(access.get(access_key, False)):
                        eligible.append(u.get('username') or uid)
        eligible.sort()
        return eligible
    except Exception as e:
        logger.warning(f"Failed to fetch eligible subadmins: {e}")
        return []

def _select_round_robin(username: str, eligible: list, counter_key: str) -> str:
    try:
        if not rtdb_available or not eligible:
            return ''
        meta_ref = fb_db.reference('users').child(username).child('meta').child(counter_key)
        idx = meta_ref.get()
        try:
            idx_int = int(idx) if idx is not None else 0
        except Exception:
            idx_int = 0
        selected = eligible[idx_int % len(eligible)]
        # increment and save back
        meta_ref.set((idx_int + 1) % max(len(eligible), 1))
        return selected
    except Exception as e:
        logger.warning(f"Failed RR selection: {e}")
        # fallback: pick first
        return eligible[0] if eligible else ''

def assign_lead_round_robin(username: str) -> str:
    """Return subadmin username to assign a lead to, using round-robin among eligible subadmins."""
    eligible = _get_eligible_subadmins(username, 'leads')
    return _select_round_robin(username, eligible, 'rr_leads_index')

def backfill_unassigned_leads(parent_username: str) -> None:
    """Assign 'assigned_to' to any existing leads missing it for the given parent user.
    This checks both the new structure ({username}/bots/{bot_id}/leads) and the old ({username}/leads).
    """
    if not rtdb_available or not parent_username:
        return
    try:
        # Determine bot_id if present
        bot_id = None
        try:
            bots_ref = fb_db.reference(f'{parent_username}/bots')
            bots = bots_ref.get() or {}
            if bots and len(bots) > 0:
                bot_id = list(bots.keys())[0]
        except Exception:
            bot_id = None
        # Helper to process a leads dict at a given base path
        def _process(base_path: str):
            try:
                lref = fb_db.reference(f'{base_path}/leads')
                lraw = lref.get() or {}
                if not isinstance(lraw, dict):
                    return
                for lead_id, l in lraw.items():
                    if not isinstance(l, dict):
                        continue
                    if l.get('assigned_to'):
                        continue
                    assignee = assign_lead_round_robin(parent_username)
                    try:
                        lref.child(lead_id).update({'assigned_to': assignee})
                    except Exception:
                        continue
            except Exception:
                pass
        # New structure first
        if bot_id:
            base = _base_path(parent_username, bot_id)
            _process(base)
        # Old structure fallback (no trailing '/leads' in base)
        _process(parent_username)
    except Exception as _:
        return

def assign_appointment_round_robin(username: str) -> str:
    """Return subadmin username to assign an appointment to, using round-robin among sub-admins with appointments access."""
    eligible = _get_eligible_subadmins(username, 'appointments')
    return _select_round_robin(username, eligible, 'rr_appts_index')

def backfill_unassigned_appointments(parent_username: str) -> None:
    if not rtdb_available or not parent_username:
        return
    try:
        bot_id = None
        try:
            bots_ref = fb_db.reference(f'{parent_username}/bots')
            bots = bots_ref.get() or {}
            if bots and len(bots) > 0:
                bot_id = list(bots.keys())[0]
        except Exception:
            bot_id = None
        def _process(base_path: str):
            try:
                aref = fb_db.reference(f'{base_path}/appointments')
                araw = aref.get() or {}
                if not isinstance(araw, dict):
                    return
                for appt_id, a in araw.items():
                    if not isinstance(a, dict):
                        continue
                    if a.get('assigned_to'):
                        continue
                    assignee = assign_appointment_round_robin(parent_username)
                    try:
                        aref.child(appt_id).update({'assigned_to': assignee})
                    except Exception:
                        continue
            except Exception:
                pass
        if bot_id:
            base = _base_path(parent_username, bot_id)
            _process(base)
        _process(parent_username)
    except Exception:
        return

def _reassign_items_round_robin(parent_username: str, removed_username: str) -> None:
    if not rtdb_available or not parent_username or not removed_username:
        return
    try:
        # Determine bot_id if present
        bot_id = None
        try:
            bots_ref = fb_db.reference(f'{parent_username}/bots')
            bots = bots_ref.get() or {}
            if bots and len(bots) > 0:
                bot_id = list(bots.keys())[0]
        except Exception:
            bot_id = None

        # Build eligible lists excluding removed user
        eligible_leads = [u for u in _get_eligible_subadmins(parent_username, 'leads') if u != removed_username]
        eligible_appts = [u for u in _get_eligible_subadmins(parent_username, 'appointments') if u != removed_username]

        def reassign_collection(base_path: str, collection: str, eligible: list, counter_start: int = 0):
            if not eligible:
                return
            try:
                ref = fb_db.reference(f'{base_path}/{collection}')
                raw = ref.get() or {}
                if not isinstance(raw, dict):
                    return
                idx = max(0, int(counter_start))
                for item_id, obj in raw.items():
                    if not isinstance(obj, dict):
                        continue
                    if obj.get('assigned_to') == removed_username:
                        assignee = eligible[idx % len(eligible)]
                        try:
                            ref.child(item_id).update({'assigned_to': assignee})
                        except Exception:
                            pass
                        idx += 1
            except Exception:
                return

        # New structure
        if bot_id:
            base = _base_path(parent_username, bot_id)
            reassign_collection(base, 'leads', eligible_leads)
            reassign_collection(base, 'appointments', eligible_appts)
        # Old structure fallbacks
        reassign_collection(parent_username, 'leads', eligible_leads)
        reassign_collection(parent_username, 'appointments', eligible_appts)
    except Exception:
        return


@app.route('/users/<username>/subadmins', methods=['DELETE'])
@require_auth
def delete_subadmin(username):
    try:
        # Verify the token subject matches the route username (only main admin can delete)
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        subject = _verify_token(token)
        if not subject or subject != username:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 401

        if not rtdb_available:
            return jsonify({'success': False, 'message': 'Database not available'}), 503

        data = request.json or {}
        sub_username = (data.get('username') or '').strip()
        if not sub_username:
            return jsonify({'success': False, 'message': 'Missing subadmin username'}), 400
        if sub_username == username:
            return jsonify({'success': False, 'message': 'Cannot delete main admin'}), 400

        # Check user exists and is child of this parent
        uref = fb_db.reference('users').child(sub_username)
        uobj = uref.get() or {}
        if not uobj or not isinstance(uobj, dict) or uobj.get('role') != 'subadmin' or uobj.get('parent') != username:
            return jsonify({'success': False, 'message': 'Subadmin not found for this user'}), 404

        # Reassign leads and appointments
        try:
            _reassign_items_round_robin(username, sub_username)
        except Exception as e:
            logger.warning(f'Reassign on delete failed: {e}')

        # Finally delete the subadmin user record
        try:
            uref.delete()
        except Exception as e:
            logger.warning(f'Failed deleting subadmin record: {e}')

        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Delete subadmin error: {e}")
        return jsonify({'success': False, 'message': 'Failed to delete subadmin'}), 500

if __name__ == '__main__':
    # Load sample knowledge base and create default user for testing
    load_sample_knowledge_base()
    create_default_user()
    app.run(debug=True, port=5001) 